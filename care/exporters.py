"""Utilities to export care records to multiple formats."""
from __future__ import annotations

import csv
import re
import textwrap
import unicodedata
import zipfile
from dataclasses import dataclass
import math
from datetime import date, datetime
from io import BytesIO
from typing import Iterable, Sequence
from xml.sax.saxutils import escape as xml_escape

from django.http import HttpResponse
from django.utils import timezone
from django.utils.text import slugify


class ExportDependencyError(RuntimeError):
    """Raised when an export format is unavailable due to missing deps."""


@dataclass(slots=True)
class ExportMetadata:
    start: date | None
    end: date | None
    period_label: str
    patient_name: str | None
    records_total: int
    group_name: str | None = None
    record_types_label: str | None = None
    patient_identifier: str | None = None
    professional_name: str | None = None
    unit_name: str | None = None

    @property
    def range_slug(self) -> str:
        start = self.start.isoformat() if self.start else "inicio"
        end = self.end.isoformat() if self.end else "agora"
        if not self.start and not self.end:
            return "todos"
        return f"{start}_{end}"

    @property
    def patient_slug(self) -> str:
        if not self.patient_name:
            return "todos"
        return slugify(self.patient_name) or "paciente"

    def describe(self) -> str:
        parts: list[str] = [f"Período: {self.period_label}"]
        if self.start and self.end:
            parts.append(f"De {self.start.strftime('%d/%m/%Y')} até {self.end.strftime('%d/%m/%Y')}")
        elif self.start and not self.end:
            parts.append(f"A partir de {self.start.strftime('%d/%m/%Y')}")
        elif not self.start and self.end:
            parts.append(f"Até {self.end.strftime('%d/%m/%Y')}")
        if self.patient_name:
            parts.append(f"Paciente: {self.patient_name}")
        if self.group_name:
            parts.append(f"Grupo: {self.group_name}")
        if self.record_types_label:
            parts.append(f"Tipos: {self.record_types_label}")
        parts.append(f"Total de registros: {self.records_total}")
        return " | ".join(parts)

    def summary_rows(self) -> list[tuple[str, str]]:
        rows: list[tuple[str, str]] = [("Período selecionado", self.period_label)]
        if self.start and self.end:
            rows.append(("Intervalo", f"{self.start.strftime('%d/%m/%Y')} – {self.end.strftime('%d/%m/%Y')}"))
        elif self.start and not self.end:
            rows.append(("Intervalo", f"Desde {self.start.strftime('%d/%m/%Y')}"))
        elif not self.start and self.end:
            rows.append(("Intervalo", f"Até {self.end.strftime('%d/%m/%Y')}"))
        rows.append(("Paciente", self.patient_name or "Todos os pacientes"))
        if self.patient_identifier:
            rows.append(("Identificador do paciente", self.patient_identifier))
        rows.append(("Grupo", self.group_name or "Todos os grupos"))
        rows.append(("Tipos de registro", self.record_types_label or "Todos os tipos"))
        rows.append(("Total de registros", str(self.records_total)))
        rows.append(("Gerado em", timezone.localtime().strftime("%d/%m/%Y %H:%M")))
        return rows


def _clean_text(value: str | None) -> str:
    return (value or "").replace("\r", " ").replace("\n", " ").strip()


COLUMNS: Sequence[tuple[str, str]] = (
    ("date", "Data"),
    ("time", "Hora"),
    ("category", "Categoria"),
    ("what", "O que"),
    ("description", "Observações"),
    ("caregiver", "Cuidador"),
    ("patient", "Paciente"),
    ("status", "Status"),
    ("exception", "Exceção?"),
)

MEDICATION_COLUMNS: Sequence[tuple[str, str]] = (
    ("created_date", "Data"),
    ("created_time", "Horário"),
    ("recurrence", "Recorrência"),
    ("medication", "Medicamento"),
    ("patient", "Paciente"),
    ("caregiver", "Cuidador"),
)

DEFAULT_COL_WIDTH_UNITS = 11
CREATED_COL_WIDTH_UNITS = 15
DEFAULT_COL_WIDTH_PX = 82
CREATED_COL_WIDTH_PX = 110
DOCUMENT_TITLE = "Relatório de registros de cuidado"
SLEEP_DOCUMENT_TITLE = "Relatório de Sono"


@dataclass(slots=True)
class SleepExportLayout:
    title: str
    generated_at: str
    patient_label: str
    professional_label: str
    period_label: str
    summary_cards: list[tuple[str, str]]
    start_count: int
    end_count: int
    total_count: int
    history_columns: list[str]
    history_rows: list[dict[str, str]]
    clinical_observations: list[str]
    technical_notes: list[str]


@dataclass(slots=True)
class BathroomExportLayout:
    title: str
    generated_at: str
    patient_label: str
    professional_label: str
    period_label: str
    summary_cards: list[tuple[str, str]]
    history_columns: list[str]
    history_rows: list[dict[str, str]]
    technical_notes: list[str]


@dataclass(slots=True)
class VitalExportLayout:
    title: str
    generated_at: str
    patient_label: str
    professional_label: str
    period_label: str
    summary_cards: list[tuple[str, str, str]]
    history_columns: list[str]
    history_rows: list[dict[str, str]]
    technical_notes: list[str]


@dataclass(slots=True)
class MedicationExportLayout:
    title: str
    generated_at: str
    patient_label: str
    professional_label: str
    period_label: str
    summary_cards: list[tuple[str, str, str]]
    history_columns: list[str]
    history_rows: list[dict[str, str]]
    technical_notes: list[str]


@dataclass(slots=True)
class ActivityExportLayout:
    title: str
    generated_at: str
    patient_label: str
    professional_label: str
    period_label: str
    summary_cards: list[tuple[str, str, str]]
    history_columns: list[str]
    history_rows: list[dict[str, str]]
    technical_notes: list[str]


@dataclass(slots=True)
class MealExportLayout:
    title: str
    generated_at: str
    patient_label: str
    professional_label: str
    period_label: str
    summary_cards: list[tuple[str, str, str]]
    history_columns: list[str]
    history_rows: list[dict[str, str]]
    technical_notes: list[str]


@dataclass(slots=True)
class ProgressExportLayout:
    title: str
    generated_at: str
    patient_label: str
    professional_label: str
    period_label: str
    summary_cards: list[tuple[str, str, str]]
    history_columns: list[str]
    history_rows: list[dict[str, str]]
    technical_notes: list[str]


@dataclass(slots=True)
class ConsolidatedExportSection:
    type_value: str
    label: str
    rows: list[dict[str, str]]
    columns: Sequence[tuple[str, str]]


def _format_decimal(value: float, places: int = 1) -> str:
    return f"{value:.{places}f}".replace(".", ",")


def _parse_export_date(value: str) -> date | None:
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except (TypeError, ValueError):
        return None


def _parse_export_datetime(row: dict[str, str]) -> datetime | None:
    row_date = _parse_export_date(row.get("date", ""))
    row_time = row.get("time", "")
    if not row_date or not row_time:
        return None
    try:
        parsed_time = datetime.strptime(row_time, "%H:%M").time()
    except ValueError:
        return None
    return datetime.combine(row_date, parsed_time)


def _human_date(value: str) -> str:
    parsed = _parse_export_date(value)
    return parsed.strftime("%d/%m/%Y") if parsed else value


def _duration_label(start_dt: datetime, end_dt: datetime) -> str:
    minutes = int(round((end_dt - start_dt).total_seconds() / 60))
    if minutes <= 0:
        return ""
    hours, remainder = divmod(minutes, 60)
    return f"{hours}h {remainder:02d}min"


def _sleep_event(row: dict[str, str]) -> str:
    return (row.get("what") or "").strip()


def _sleep_observation(*records: dict[str, str]) -> str:
    parts: list[str] = []
    for row in records:
        description = (row.get("description") or "").strip()
        if description:
            parts.append(description)
    return " | ".join(parts)


def _status_key(value: str) -> str:
    normalized = (value or "").strip().lower()
    if normalized in {"realizada", "realizado", "done"}:
        return "done"
    if normalized in {"pendente", "pending"}:
        return "pending"
    if normalized in {"não realizado", "nao realizado", "missed"}:
        return "missed"
    return ""


def _status_label(value: str) -> str:
    key = _status_key(value)
    if key == "done":
        return "REALIZADA"
    if key == "pending":
        return "PENDENTE"
    if key == "missed":
        return "NÃO REALIZADO"
    return (value or "").strip().upper()


def is_sleep_export(rows: list[dict[str, str]], meta: ExportMetadata) -> bool:
    if meta.record_types_label == "Sono":
        return True
    return bool(rows) and all((row.get("category") or "").strip() == "Sono" for row in rows)


def is_bathroom_export(rows: list[dict[str, str]], meta: ExportMetadata) -> bool:
    if meta.record_types_label == "Banheiro":
        return True
    return bool(rows) and all((row.get("category") or "").strip() == "Banheiro" for row in rows)


def _normalize_bathroom_label(value: str) -> str:
    raw = (value or "").strip()
    lowered = raw.lower()
    lowered = (
        lowered.replace("ã", "a")
        .replace("á", "a")
        .replace("â", "a")
        .replace("ç", "c")
        .replace("é", "e")
        .replace("ê", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ô", "o")
        .replace("õ", "o")
        .replace("ú", "u")
    )
    if "sem ocorr" in lowered:
        return "Sem ocorrência durante o dia"
    if "higien" in lowered and "oral" in lowered:
        return "Higienização Oral"
    if "banho" in lowered:
        return "Banho"
    if "urin" in lowered:
        return "Urina"
    if "evac" in lowered:
        return "Evacuação"
    if "vomit" in lowered:
        return "Vômito"
    return raw or "Outro"


def _bathroom_tag_kind(label: str) -> str:
    normalized = _normalize_bathroom_label(label)
    if normalized in {"Higienização Oral", "Banho"}:
        return "hygiene"
    if normalized in {"Urina", "Evacuação"}:
        return "elimination"
    if normalized == "Vômito":
        return "alert"
    return "neutral"


def build_bathroom_export_layout(rows: list[dict[str, str]], meta: ExportMetadata) -> BathroomExportLayout:
    ordered_rows = sorted(rows, key=lambda row: (row.get("date", ""), row.get("time", ""), row.get("patient", "")))
    total = len(ordered_rows)
    hygiene_count = 0
    elimination_count = 0
    alert_count = 0
    history_rows: list[dict[str, str]] = []

    for row in ordered_rows:
        occurrence = _normalize_bathroom_label(row.get("what", ""))
        tag_kind = _bathroom_tag_kind(occurrence)
        if tag_kind == "hygiene":
            hygiene_count += 1
        elif tag_kind == "elimination":
            elimination_count += 1
        elif tag_kind == "alert":
            alert_count += 1
        history_rows.append(
            {
                "date": _human_date(row.get("date", "")) or "Não informado",
                "time": row.get("time", "") or "Não informado",
                "occurrence": occurrence,
                "tag_kind": tag_kind,
                "recurrence": row.get("recurrence", "") or "Não informado",
                "observation": row.get("description", "").strip(),
            }
        )

    return BathroomExportLayout(
        title="Relatório de Banheiro e Higiene",
        generated_at=timezone.localtime().strftime("%d/%m/%Y %H:%M"),
        patient_label=meta.patient_name or "Todos os pacientes",
        professional_label=meta.professional_name or "Não informado",
        period_label=meta.period_label or "Todos os tempos",
        summary_cards=[
            ("Total de Registros", str(total)),
            ("Higiene & Banho", str(hygiene_count)),
            ("Eliminações (Urina/Evac.)", str(elimination_count)),
            ("Intercorrências (Vômito)", str(alert_count)),
        ],
        history_columns=["DATA", "HORÁRIO", "TIPO DE OCORRÊNCIA", "RECORRÊNCIA", "OBSERVAÇÕES"],
        history_rows=history_rows,
        technical_notes=[
            "Tipos Monitorados: Urina, Evacuação, Banho, Vômito, Higienização Oral, Outro, Sem ocorrência durante o dia.",
            "Recorrência: Define o padrão temporal observado para a atividade ou sintoma no prontuário do paciente.",
            "Alertas e episódios de refluxo ou vômito devem ser imediatamente destacados e informados à equipe clínica.",
        ],
    )


def is_vital_export(rows: list[dict[str, str]], meta: ExportMetadata) -> bool:
    if meta.record_types_label == "Sinais Vitais":
        return True
    return bool(rows) and all((row.get("category") or "").strip() == "Sinais Vitais" for row in rows)


def _normalize_vital_text(value: str) -> str:
    normalized = (value or "").strip().lower()
    normalized = unicodedata.normalize("NFKD", normalized)
    return "".join(char for char in normalized if not unicodedata.combining(char))
    return (
        normalized.replace("ã", "a")
        .replace("á", "a")
        .replace("â", "a")
        .replace("ç", "c")
        .replace("é", "e")
        .replace("ê", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ô", "o")
        .replace("ú", "u")
    )


def _split_vital_what(value: str) -> tuple[str, str]:
    raw = (value or "").strip()
    if not raw:
        return "Não informado", "Não informado"
    for separator in (" • ", " â€¢ ", " - ", " | ", " / "):
        if separator in raw:
            left, right = raw.split(separator, 1)
            return left.strip() or "Não informado", right.strip() or "Não informado"
    return raw, "Não informado"


def _vital_status_level(status: str) -> str:
    normalized = _normalize_vital_text(status)
    if "normal" in normalized:
        return "normal"
    critical_terms = (
        "baixa saturacao",
        "saturacao baixa",
        "hipertensao grave",
        "hipertenso grave",
        "hipotensao grave",
        "hipotenso grave",
        "hipotermia severa",
        "critico",
        "critica",
        "grave",
    )
    if any(term in normalized for term in critical_terms):
        return "alert"
    warning_terms = (
        "febre",
        "taquicardia",
        "hipertenso",
        "hipertens",
        "hipotenso",
        "hipotens",
        "bradicardia",
        "hipotermia",
        "alterado",
        "atencao",
    )
    if any(term in normalized for term in warning_terms):
        return "warning"
    return "warning" if normalized and normalized != "nao informado" else "neutral"


def build_vital_export_layout(rows: list[dict[str, str]], meta: ExportMetadata) -> VitalExportLayout:
    ordered_rows = sorted(rows, key=lambda row: (row.get("date", ""), row.get("time", ""), row.get("patient", "")))
    normal_count = 0
    warning_count = 0
    alert_count = 0
    history_rows: list[dict[str, str]] = []

    for row in ordered_rows:
        vital_name, clinical_status = _split_vital_what(row.get("what", ""))
        status_level = _vital_status_level(clinical_status)
        if status_level == "normal":
            normal_count += 1
        elif status_level == "alert":
            alert_count += 1
        elif status_level == "warning":
            warning_count += 1
        history_rows.append(
            {
                "date": _human_date(row.get("date", "")) or "Não informado",
                "time": row.get("time", "") or "Não informado",
                "vital": vital_name,
                "recurrence": row.get("recurrence", "") or "Não informado",
                "status": clinical_status,
                "status_level": status_level,
                "observation": row.get("description", "").strip(),
            }
        )

    return VitalExportLayout(
        title="Relatório de Sinais Vitais",
        generated_at=timezone.localtime().strftime("%d/%m/%Y %H:%M"),
        patient_label=meta.patient_name or "Todos os pacientes",
        professional_label=meta.professional_name or "Não informado",
        period_label=meta.period_label or "Todos os tempos",
        summary_cards=[
            ("Total Medições", str(len(ordered_rows)), "main"),
            ("Status Normal", str(normal_count), "normal"),
            ("Atenção (Febre/Taquic.)", str(warning_count), "warning"),
            ("Alerta Crítico", str(alert_count), "alert"),
        ],
        history_columns=["DATA", "HORÁRIO", "SINAL VITAL (QUAL)", "RECORRÊNCIA", "STATUS CLÍNICO", "OBSERVAÇÕES"],
        history_rows=history_rows,
        technical_notes=[
            "Sinais Monitorados: Pressão Arterial (PA), Frequência Cardíaca (FrC), SpO2 (Oxímetro), Temperatura, Outro.",
            "Classificação de Status: Normal, Hipertenso, Hipotenso, Febre, Hipotermia, Taquicardia, Bradicardia, Baixa Saturação, Outro.",
            "Parâmetros em estado de Alerta Crítico ou que apresentem refratariedade à medicação exigem notificação médica imediata.",
        ],
    )


def is_medication_export(rows: list[dict[str, str]], meta: ExportMetadata) -> bool:
    if meta.record_types_label == "Medicação":
        return True
    return bool(rows) and all((row.get("category") or "").strip() == "Medicação" for row in rows)


def _split_medication_name_dose(value: str) -> tuple[str, str]:
    raw = (value or "").strip()
    if not raw:
        return "Não informado", ""
    for label in ("Remédio/Dose:", "Remédio:", "Medicamento:"):
        if raw.startswith(label):
            raw = raw[len(label):].strip()
    for separator in (" • ", " â€¢ ", " - ", " | ", " / "):
        if separator in raw:
            left, right = raw.split(separator, 1)
            left = left.strip()
            right = right.strip()
            if left and right:
                return left, right
    match = re.match(r"^(.+?)\s+(\d[\w\s,./%µ]+)$", raw)
    if match:
        return match.group(1).strip(), match.group(2).strip()
    return raw, ""


def _medication_unit_from_text(*values: str) -> str:
    text = _normalize_vital_text(" ".join(value or "" for value in values))
    if "gota" in text:
        return "Gotas"
    if "comprim" in text:
        return "Comprimidos"
    if "capsul" in text:
        return "Cápsulas"
    if re.search(r"\bml\b", text):
        return "mL"
    return "Cápsulas"


def _quantity_badge(quantity: str, unit: str) -> str:
    qty = (quantity or "").strip()
    if not qty:
        return "Não informado"
    try:
        qty_int = int(float(qty.replace(",", ".")))
    except ValueError:
        return qty
    if unit == "Gotas":
        return f"{qty_int} Gota" if qty_int == 1 else f"{qty_int} Gotas"
    if unit == "Comprimidos":
        return f"{qty_int} Comprimido" if qty_int == 1 else f"{qty_int} Comprimidos"
    if unit == "mL":
        return f"{qty_int} mL"
    return f"{qty_int} Cápsula" if qty_int == 1 else f"{qty_int} Cápsulas"


def _has_medication_intercurrence(row: dict[str, str]) -> bool:
    text = _normalize_vital_text(
        " ".join(
            [
                row.get("status", ""),
                row.get("exception", ""),
                row.get("description", ""),
                row.get("missed_reason", ""),
            ]
        )
    )
    negative_terms = (
        "sem reacao",
        "sem reacoes",
        "sem reacao adversa",
        "sem reacoes adversas",
        "nenhuma reacao",
        "nenhuma intercorrencia",
        "sem intercorrencia",
    )
    if any(term in text for term in negative_terms):
        text = text.replace("sem reacoes adversas", "")
        text = text.replace("sem reacao adversa", "")
        text = text.replace("sem reacoes", "")
        text = text.replace("sem reacao", "")
        text = text.replace("nenhuma reacao", "")
        text = text.replace("nenhuma intercorrencia", "")
        text = text.replace("sem intercorrencia", "")
    terms = (
        "recusa",
        "recusou",
        "esquec",
        "nao realizado",
        "não realizado",
        "intercorr",
        "reacao",
        "reacao adversa",
        "adversa",
        "vomit",
        "alerg",
    )
    return row.get("exception") == "Sim" or any(term in text for term in terms)


def build_medication_export_layout(rows: list[dict[str, str]], meta: ExportMetadata) -> MedicationExportLayout:
    ordered_rows = sorted(rows, key=lambda row: (row.get("date", row.get("created_date", "")), row.get("time", row.get("created_time", "")), row.get("patient", "")))
    capsules_count = 0
    drops_count = 0
    intercurrence_count = 0
    history_rows: list[dict[str, str]] = []

    for row in ordered_rows:
        med_name = row.get("medication_name") or ""
        dose = row.get("dose") or ""
        if not med_name:
            med_name, parsed_dose = _split_medication_name_dose(row.get("medication") or row.get("what", ""))
            dose = dose or parsed_dose
        unit = _medication_unit_from_text(row.get("quantity_unit", ""), med_name, dose, row.get("description", ""))
        if unit in {"Cápsulas", "Comprimidos"}:
            capsules_count += 1
        elif unit == "Gotas":
            drops_count += 1
        if _has_medication_intercurrence(row):
            intercurrence_count += 1
        history_rows.append(
            {
                "date": _human_date(row.get("date", row.get("created_date", ""))) or "Não informado",
                "time": row.get("time", row.get("created_time", "")) or "Não informado",
                "medication_name": med_name or "Não informado",
                "dose": dose,
                "quantity": _quantity_badge(row.get("quantity", ""), unit),
                "recurrence": row.get("recurrence", "") or "Não informado",
                "observation": row.get("description", "").strip(),
            }
        )

    return MedicationExportLayout(
        title="Relatório de Medicação",
        generated_at=timezone.localtime().strftime("%d/%m/%Y %H:%M"),
        patient_label=meta.patient_name or "Todos os pacientes",
        professional_label=meta.professional_name or "Não informado",
        period_label=meta.period_label or "Todos os tempos",
        summary_cards=[
            ("Total de Ministrações", str(len(ordered_rows)), "main"),
            ("Via Oral (Cápsulas)", str(capsules_count), "blue"),
            ("Via Oral (Gotas)", str(drops_count), "blue"),
            ("Intercorrências / Recusas", str(intercurrence_count), "neutral"),
        ],
        history_columns=["DATA", "HORÁRIO", "MEDICAMENTO / DOSE", "QUANTIDADE", "RECORRÊNCIA", "OBSERVAÇÕES"],
        history_rows=history_rows,
        technical_notes=[
            "Campos Obrigatórios: Toda administração exige o registro exato de Data, Horário, Nome do Medicamento, Dose e Quantidade Utilizada.",
            "Atenção à Forma Farmacêutica: Diferenciar rigorosamente a unidade de medida entre Cápsulas, Comprimidos, Gotas, mL e outras formas farmacêuticas para evitar erros de administração.",
            "Qualquer sintoma atípico, recusa do paciente ou esquecimento de dose deve ser detalhado imediatamente no campo de observações.",
        ],
    )


def is_activity_export(rows: list[dict[str, str]], meta: ExportMetadata) -> bool:
    if meta.record_types_label == "Atividade":
        return True
    return bool(rows) and all((row.get("category") or "").strip() == "Atividade" for row in rows)


def _activity_text(*values: str) -> str:
    return _normalize_vital_text(" ".join(value or "" for value in values))


def _is_cardiovascular_activity(name: str) -> bool:
    normalized = _activity_text(name)
    terms = (
        "caminhada",
        "bicicleta ergometrica",
        "bicicleta",
        "exercicio aerobico",
        "exercicios aerobicos",
        "aerobico",
        "cardiovascular",
        "esteira",
        "corrida",
    )
    return any(term in normalized for term in terms)


def _is_mobility_activity(name: str) -> bool:
    normalized = _activity_text(name)
    terms = (
        "alongamento",
        "mobilidade articular",
        "mobilidade",
        "flexibilidade",
        "fisioterapia motora",
        "fisioterapia",
        "reabilitacao funcional",
        "reabilitacao",
    )
    return any(term in normalized for term in terms)


def _has_activity_refusal_or_suspension(row: dict[str, str]) -> bool:
    text = _activity_text(
        row.get("what", ""),
        row.get("status", ""),
        row.get("description", ""),
        row.get("missed_reason", ""),
    )
    terms = (
        "recusa",
        "recusou",
        "recusada",
        "recusado",
        "interrompida",
        "interrompido",
        "suspensa",
        "suspenso",
        "suspensao",
        "impossibilidade",
        "impossivel",
        "nao realizado",
        "nao realizada",
        "cancelada",
        "cancelado",
    )
    return row.get("exception") == "Sim" or any(term in text for term in terms)


def build_activity_export_layout(rows: list[dict[str, str]], meta: ExportMetadata) -> ActivityExportLayout:
    ordered_rows = sorted(rows, key=lambda row: (row.get("date", ""), row.get("time", ""), row.get("patient", "")))
    cardiovascular_count = 0
    mobility_count = 0
    refusal_count = 0
    history_rows: list[dict[str, str]] = []

    for row in ordered_rows:
        activity_name = (row.get("what") or "").strip() or "Não informado"
        if _is_cardiovascular_activity(activity_name):
            cardiovascular_count += 1
        if _is_mobility_activity(activity_name):
            mobility_count += 1
        if _has_activity_refusal_or_suspension(row):
            refusal_count += 1
        history_rows.append(
            {
                "date": _human_date(row.get("date", "")) or "Não informado",
                "time": row.get("time", "") or "Não informado",
                "activity_name": activity_name,
                "recurrence": row.get("recurrence", "") or "Não informado",
                "observation": (row.get("description") or "").strip(),
            }
        )

    return ActivityExportLayout(
        title="Relatório de Atividade Física",
        generated_at=timezone.localtime().strftime("%d/%m/%Y %H:%M"),
        patient_label=meta.patient_name or "Todos os pacientes",
        professional_label=meta.professional_name or "Não informado",
        period_label=meta.period_label or "Todos os tempos",
        summary_cards=[
            ("Total de Exercícios", str(len(ordered_rows)), "main"),
            ("Cardiovascular", str(cardiovascular_count), "blue"),
            ("Mobilidade/Alongamento", str(mobility_count), "blue"),
            ("Recusas ou Suspensões", str(refusal_count), "neutral"),
        ],
        history_columns=["DATA", "HORÁRIO", "ATIVIDADE (NOME)", "RECORRÊNCIA", "OBSERVAÇÕES"],
        history_rows=history_rows,
        technical_notes=[
            "Critérios de Registro: Toda entrada deve documentar com precisão o Nome da Atividade, Horário de início e a Recorrência prevista.",
            "Monitoramento Preventivo: Recomenda-se aferir os sinais vitais antes e após atividades de maior intensidade cardiovascular.",
            "Sinais de dor persistente, tontura, falta de ar crônica ou fadiga extrema exigem a interrupção imediata da atividade e anotação detalhada.",
        ],
    )


def is_meal_export(rows: list[dict[str, str]], meta: ExportMetadata) -> bool:
    if meta.record_types_label == "Alimentação":
        return True
    return bool(rows) and all((row.get("category") or "").strip() == "Alimentação" for row in rows)


def _split_meal_what(value: str) -> tuple[str, str]:
    raw = (value or "").strip()
    if not raw:
        return "Não informado", "Outro"
    for separator in (" • ", " â€¢ ", " Ã¢â‚¬Â¢ ", " - ", " | ", " / "):
        if separator in raw:
            left, right = raw.split(separator, 1)
            return left.strip() or "Não informado", right.strip() or "Outro"
    normalized = _normalize_vital_text(raw)
    if "boa aceitacao" in normalized:
        return raw, "Boa aceitação"
    if "ruim aceitacao" in normalized or "baixa aceitacao" in normalized:
        return raw, "Ruim aceitação"
    return raw, "Outro"


def _meal_acceptance_kind(value: str) -> str:
    normalized = _normalize_vital_text(value)
    if "boa aceitacao" in normalized or normalized in {"boa", "aceitou bem"}:
        return "good"
    if "ruim aceitacao" in normalized or "baixa aceitacao" in normalized or normalized in {"ruim", "pouca aceitacao"}:
        return "poor"
    return "other"


def _meal_acceptance_label(value: str) -> str:
    kind = _meal_acceptance_kind(value)
    if kind == "good":
        return "Boa Aceitação"
    if kind == "poor":
        return "Ruim Aceitação"
    raw = (value or "").strip()
    return raw if raw and raw.lower() != "outro" else "Outro"


def build_meal_export_layout(rows: list[dict[str, str]], meta: ExportMetadata) -> MealExportLayout:
    ordered_rows = sorted(rows, key=lambda row: (row.get("date", ""), row.get("time", ""), row.get("patient", "")))
    good_count = 0
    poor_count = 0
    other_count = 0
    history_rows: list[dict[str, str]] = []

    for row in ordered_rows:
        meal_name, acceptance_raw = _split_meal_what(row.get("what", ""))
        acceptance_kind = _meal_acceptance_kind(acceptance_raw)
        if acceptance_kind == "good":
            good_count += 1
        elif acceptance_kind == "poor":
            poor_count += 1
        else:
            other_count += 1
        history_rows.append(
            {
                "date": _human_date(row.get("date", "")) or "Não informado",
                "time": row.get("time", "") or "Não informado",
                "meal_name": meal_name or "Não informado",
                "recurrence": row.get("recurrence", "") or "Não informado",
                "acceptance": _meal_acceptance_label(acceptance_raw),
                "acceptance_kind": acceptance_kind,
                "observation": (row.get("description") or "").strip(),
            }
        )

    return MealExportLayout(
        title="Relatório de Alimentação",
        generated_at=timezone.localtime().strftime("%d/%m/%Y %H:%M"),
        patient_label=meta.patient_name or "Todos os pacientes",
        professional_label=meta.professional_name or "Não informado",
        period_label=meta.period_label or "Todos os tempos",
        summary_cards=[
            ("Refeições Registradas", str(len(ordered_rows)), "main"),
            ("Boa Aceitação", str(good_count), "good"),
            ("Ruim Aceitação", str(poor_count), "poor"),
            ("Outras Ocorrências", str(other_count), "blue"),
        ],
        history_columns=["DATA", "HORÁRIO", "REFEIÇÃO", "RECORRÊNCIA", "ACEITAÇÃO", "OBSERVAÇÕES"],
        history_rows=history_rows,
        technical_notes=[
            "Padrões de Refeição: Café da Manhã, Lanche da Manhã, Almoço, Lanche da Tarde, Jantar, Ceia da noite, Outro.",
            "Grau de Aceitação: Mapeado rigorosamente entre Boa Aceitação, Ruim Aceitação e Outro.",
            "Padrões repetidos de Ruim Aceitação devem ser reportados à equipe de nutrição para ajuste da consistência ou do cardápio.",
        ],
    )


def is_progress_export(rows: list[dict[str, str]], meta: ExportMetadata) -> bool:
    if meta.record_types_label == "Evolução/Regressão":
        return True
    return bool(rows) and all((row.get("category") or "").strip() == "Evolução/Regressão" for row in rows)


def _progress_kind(value: str) -> str:
    normalized = _normalize_vital_text(value)
    if "evolucao" in normalized or normalized in {"evolution", "melhora", "progresso"}:
        return "evolution"
    if "regressao" in normalized or normalized in {"regression", "piora"}:
        return "regression"
    return "other"


def _progress_label(value: str) -> str:
    kind = _progress_kind(value)
    if kind == "evolution":
        return "Evolução"
    if kind == "regression":
        return "Regressão"
    raw = (value or "").strip()
    return raw if raw and raw.lower() != "outro" else "Outro"


def build_progress_export_layout(rows: list[dict[str, str]], meta: ExportMetadata) -> ProgressExportLayout:
    ordered_rows = sorted(rows, key=lambda row: (row.get("date", ""), row.get("time", ""), row.get("patient", "")))
    evolution_count = 0
    regression_count = 0
    other_count = 0
    history_rows: list[dict[str, str]] = []

    for row in ordered_rows:
        classification_raw = row.get("progress_trend") or row.get("what") or ""
        classification_kind = _progress_kind(classification_raw)
        if classification_kind == "evolution":
            evolution_count += 1
        elif classification_kind == "regression":
            regression_count += 1
        else:
            other_count += 1
        history_rows.append(
            {
                "date": _human_date(row.get("date", "")) or "Não informado",
                "time": row.get("time", "") or "Não informado",
                "classification": _progress_label(classification_raw),
                "classification_kind": classification_kind,
                "recurrence": row.get("recurrence", "") or "Não informado",
                "observation": (row.get("description") or "").strip(),
            }
        )

    return ProgressExportLayout(
        title="Relatório de Evolução / Regressão",
        generated_at=timezone.localtime().strftime("%d/%m/%Y %H:%M"),
        patient_label=meta.patient_name or "Todos os pacientes",
        professional_label=meta.professional_name or "Não informado",
        period_label=meta.period_label or "Todos os tempos",
        summary_cards=[
            ("Total de Registros", str(len(ordered_rows)), "main"),
            ("Eventos de Evolução", str(evolution_count), "evolution"),
            ("Eventos de Regressão", str(regression_count), "regression"),
            ("Outras Ocorrências", str(other_count), "neutral"),
        ],
        history_columns=["DATA", "HORÁRIO", "CLASSIFICAÇÃO", "RECORRÊNCIA", "OBSERVAÇÕES / NOTAS CLÍNICAS"],
        history_rows=history_rows,
        technical_notes=[
            "Classificações Clínicas: Registradas estritamente sob as categorias de Evolução, Regressão ou Outro.",
            "Critério de Análise: Toda mudança cognitiva, comportamental, funcional ou motora marcante deve ser reportada com detalhamento cronológico de Horário e Recorrência.",
            "Apontamentos reiterados de Regressão devem disparar um alerta para reavaliação da conduta ou abordagem terapêutica atual.",
        ],
    )


def _build_sleep_export_layout_legacy(rows: list[dict[str, str]], meta: ExportMetadata):
    header_rows: list[tuple[str, str]] = [
        ("Título do documento", SLEEP_DOCUMENT_TITLE),
        ("Paciente", meta.patient_name or "Todos os pacientes"),
        ("Identificador do paciente", meta.patient_identifier or ""),
        ("Data de geração", timezone.localtime().strftime("%d/%m/%Y %H:%M")),
        ("Período analisado", meta.period_label),
        ("Profissional responsável", meta.professional_name or ""),
        ("Unidade/estabelecimento", meta.unit_name or meta.group_name or ""),
        ("Total de registros de sono", str(meta.records_total)),
    ]
    if meta.start and meta.end:
        header_rows.append(("Intervalo", f"{meta.start.strftime('%d/%m/%Y')} a {meta.end.strftime('%d/%m/%Y')}"))
    elif meta.start:
        header_rows.append(("Intervalo", f"Desde {meta.start.strftime('%d/%m/%Y')}"))
    elif meta.end:
        header_rows.append(("Intervalo", f"Até {meta.end.strftime('%d/%m/%Y')}"))
    header_rows = [(label, value) for label, value in header_rows if value]

    ordered_rows = sorted(
        rows,
        key=lambda row: (
            row.get("patient", ""),
            row.get("date", ""),
            row.get("time", ""),
        ),
    )
    include_patient = len({row.get("patient", "") for row in ordered_rows if row.get("patient")}) > 1
    pending_sleep: dict[str, dict[str, str]] = {}
    sessions: list[tuple[dict[str, str], dict[str, str], datetime, datetime]] = []
    unpaired: list[dict[str, str]] = []

    for row in ordered_rows:
        patient = row.get("patient") or ""
        event = _sleep_event(row).lower()
        row_dt = _parse_export_datetime(row)
        if not row_dt:
            unpaired.append(row)
            continue
        if event == "dormiu":
            previous = pending_sleep.get(patient)
            if previous:
                unpaired.append(previous)
            pending_sleep[patient] = row
            continue
        if event == "acordou":
            start_row = pending_sleep.pop(patient, None)
            start_dt = _parse_export_datetime(start_row or {})
            if start_row and start_dt and row_dt > start_dt:
                sessions.append((start_row, row, start_dt, row_dt))
            else:
                unpaired.append(row)
            continue
        unpaired.append(row)

    unpaired.extend(pending_sleep.values())
    durations = [(end_dt - start_dt).total_seconds() / 3600 for _, _, start_dt, end_dt in sessions]

    summary_rows: list[tuple[str, str]] = []
    if durations:
        summary_rows.append(("Média de horas dormidas", f"{_format_decimal(sum(durations) / len(durations), 1)} h"))
        summary_rows.append(("Total de períodos completos", str(len(durations))))
        summary_rows.append(("Maior duração registrada", f"{_format_decimal(max(durations), 1)} h"))
        summary_rows.append(("Menor duração registrada", f"{_format_decimal(min(durations), 1)} h"))
    summary_rows.append(("Registros de início do sono", str(sum(1 for row in ordered_rows if _sleep_event(row).lower() == "dormiu"))))
    summary_rows.append(("Registros de término/despertar", str(sum(1 for row in ordered_rows if _sleep_event(row).lower() == "acordou"))))
    summary_rows.append(("Registros avaliados", str(len(ordered_rows))))

    history_columns = ["Data"]
    if include_patient:
        history_columns.append("Paciente")
    history_columns.extend(["Início do sono", "Término", "Duração total", "Status/Responsável", "Observações"])
    history_rows: list[list[str]] = []
    for start_row, end_row, start_dt, end_dt in sessions:
        status_parts = []
        status = (end_row.get("status") or start_row.get("status") or "").strip()
        caregiver = (end_row.get("caregiver") or start_row.get("caregiver") or "").strip()
        if status:
            status_parts.append(status)
        if caregiver:
            status_parts.append(caregiver)
        row = [start_dt.strftime("%d/%m/%Y")]
        if include_patient:
            row.append(start_row.get("patient", ""))
        row.extend(
            [
                start_dt.strftime("%H:%M"),
                end_dt.strftime("%H:%M"),
                _duration_label(start_dt, end_dt),
                " - ".join(status_parts),
                _sleep_observation(start_row, end_row),
            ]
        )
        history_rows.append(row)

    for row_data in sorted(unpaired, key=lambda row: (row.get("date", ""), row.get("time", ""))):
        row = [_human_date(row_data.get("date", ""))]
        if include_patient:
            row.append(row_data.get("patient", ""))
        event = _sleep_event(row_data)
        start_time = row_data.get("time", "") if event.lower() == "dormiu" else ""
        end_time = row_data.get("time", "") if event.lower() == "acordou" else ""
        status_parts = [part for part in [row_data.get("status", ""), row_data.get("caregiver", "")] if part]
        row.extend(
            [
                start_time,
                end_time,
                "Registro incompleto",
                " - ".join(status_parts),
                _sleep_observation(row_data),
            ]
        )
        history_rows.append(row)

    clinical_observations = [
        row.get("description", "").strip()
        for row in ordered_rows
        if row.get("description", "").strip()
    ]
    technical_notes = []
    if durations:
        technical_notes.append("Duração total: diferença entre o registro Dormiu e o registro Acordou do mesmo paciente.")
        technical_notes.append("Média de horas dormidas: soma das durações completas dividida pela quantidade de períodos completos.")
    if unpaired:
        technical_notes.append("Registros sem par Dormiu/Acordou são mantidos no histórico como registro incompleto.")

    return SleepExportLayout(
        title=SLEEP_DOCUMENT_TITLE,
        header_rows=header_rows,
        summary_rows=summary_rows,
        history_columns=history_columns,
        history_rows=history_rows,
        clinical_observations=clinical_observations,
        technical_notes=technical_notes,
    )


def build_sleep_export_layout(rows: list[dict[str, str]], meta: ExportMetadata) -> SleepExportLayout:
    ordered_rows = sorted(
        rows,
        key=lambda row: (
            row.get("patient", ""),
            row.get("date", ""),
            row.get("time", ""),
        ),
    )
    include_patient_in_observation = len(
        {row.get("patient", "") for row in ordered_rows if row.get("patient")}
    ) > 1
    pending_sleep: dict[str, dict[str, str]] = {}
    sessions: list[tuple[dict[str, str], dict[str, str], datetime, datetime]] = []
    unpaired: list[dict[str, str]] = []

    for row in ordered_rows:
        patient = row.get("patient") or ""
        event = _sleep_event(row).lower()
        row_dt = _parse_export_datetime(row)
        if not row_dt:
            unpaired.append(row)
            continue
        if event == "dormiu":
            previous = pending_sleep.get(patient)
            if previous:
                unpaired.append(previous)
            pending_sleep[patient] = row
            continue
        if event == "acordou":
            start_row = pending_sleep.pop(patient, None)
            start_dt = _parse_export_datetime(start_row or {})
            if start_row and start_dt and row_dt > start_dt:
                sessions.append((start_row, row, start_dt, row_dt))
            else:
                unpaired.append(row)
            continue
        unpaired.append(row)

    unpaired.extend(pending_sleep.values())
    durations = [(end_dt - start_dt).total_seconds() / 3600 for _, _, start_dt, end_dt in sessions]
    start_count = sum(1 for row in ordered_rows if _sleep_event(row).lower() == "dormiu")
    end_count = sum(1 for row in ordered_rows if _sleep_event(row).lower() == "acordou")
    summary_cards = [
        ("Média Horas Dormidas", f"{_format_decimal(sum(durations) / len(durations), 1)} h" if durations else "Não informado"),
        ("Períodos Completos", str(len(durations))),
        ("Maior Duração", f"{_format_decimal(max(durations), 1)} h" if durations else "Não informado"),
        ("Menor Duração", f"{_format_decimal(min(durations), 1)} h" if durations else "Não informado"),
    ]

    history_rows: list[dict[str, str]] = []
    for start_row, end_row, start_dt, end_dt in sessions:
        status = (end_row.get("status") or start_row.get("status") or "").strip()
        caregiver = (end_row.get("caregiver") or start_row.get("caregiver") or "").strip()
        observation = _sleep_observation(start_row, end_row)
        if include_patient_in_observation and start_row.get("patient"):
            patient_prefix = f"Paciente: {start_row.get('patient')}"
            observation = f"{patient_prefix} | {observation}" if observation else patient_prefix
        history_rows.append(
            {
                "date": start_dt.strftime("%d/%m/%Y"),
                "start": start_dt.strftime("%H:%M"),
                "end": end_dt.strftime("%H:%M"),
                "duration": _duration_label(start_dt, end_dt),
                "duration_complete": "1",
                "status": _status_label(status),
                "status_key": _status_key(status),
                "caregiver": caregiver,
                "observation": observation,
            }
        )

    for row_data in sorted(unpaired, key=lambda row: (row.get("date", ""), row.get("time", ""))):
        event = _sleep_event(row_data).lower()
        observation = _sleep_observation(row_data)
        if include_patient_in_observation and row_data.get("patient"):
            patient_prefix = f"Paciente: {row_data.get('patient')}"
            observation = f"{patient_prefix} | {observation}" if observation else patient_prefix
        history_rows.append(
            {
                "date": _human_date(row_data.get("date", "")) or "Não informado",
                "start": row_data.get("time", "") if event == "dormiu" else "",
                "end": row_data.get("time", "") if event == "acordou" else "",
                "duration": "Registro incompleto",
                "duration_complete": "",
                "status": _status_label(row_data.get("status", "")),
                "status_key": _status_key(row_data.get("status", "")),
                "caregiver": row_data.get("caregiver", "").strip(),
                "observation": observation,
            }
        )

    clinical_observations = [
        row.get("description", "").strip()
        for row in ordered_rows
        if row.get("description", "").strip()
    ]

    return SleepExportLayout(
        title=SLEEP_DOCUMENT_TITLE,
        generated_at=timezone.localtime().strftime("%d/%m/%Y %H:%M"),
        patient_label=meta.patient_name or "Todos os pacientes",
        professional_label=meta.professional_name or "Não informado",
        period_label=meta.period_label or "Todos os tempos",
        summary_cards=summary_cards,
        start_count=start_count,
        end_count=end_count,
        total_count=len(ordered_rows),
        history_columns=["DATA", "INÍCIO", "TÉRMINO", "DURAÇÃO TOTAL", "STATUS / RESPONSÁVEL", "OBSERVAÇÕES"],
        history_rows=history_rows,
        clinical_observations=clinical_observations,
        technical_notes=[
            "Duração total: diferença entre o registro Dormiu e o registro Acordou do mesmo paciente.",
            "Média de horas dormidas: soma das durações completas dividida pela quantidade de períodos completos.",
            "Registros sem par Dormiu/Acordou são mantidos no histórico como registro incompleto.",
        ],
    )


def _column_width_units(columns: Sequence[tuple[str, str]]) -> list[int]:
    widths: list[int] = []
    for _, label in columns:
        if label == "Criado em":
            widths.append(CREATED_COL_WIDTH_UNITS)
        else:
            widths.append(DEFAULT_COL_WIDTH_UNITS)
    return widths


def _column_width_px(columns: Sequence[tuple[str, str]]) -> list[int]:
    widths: list[int] = []
    for _, label in columns:
        if label == "Criado em":
            widths.append(CREATED_COL_WIDTH_PX)
        else:
            widths.append(DEFAULT_COL_WIDTH_PX)
    return widths


def serialize_records(records: Iterable) -> list[dict[str, str]]:
    """Return a normalized list of dictionaries ready for export."""
    serialized: list[dict[str, str]] = []
    for record in records:
        serialized.append(
            {
                "date": record.date.strftime("%Y-%m-%d"),
                "time": record.time.strftime("%H:%M") if record.time else "",
                "category": record.get_type_display(),
                "what": _clean_text(record.what),
                "description": _clean_text(record.description),
                "caregiver": _clean_text(record.caregiver),
                "patient": str(record.patient),
                "status": record.get_status_display(),
                "exception": "Sim" if record.is_exception else "Não",
            }
        )
    return serialized


def serialize_medication_export(records: Iterable, patient_name: str | None) -> list[dict[str, str]]:
    """Return rows for the medication export table."""
    serialized: list[dict[str, str]] = []
    for record in records:
        created_at = timezone.localtime(record.timestamp)
        medication_name = ""
        dose = ""
        if record.medication:
            medication_name = _clean_text(record.medication.name)
            dose = _clean_text(record.medication.dosage)
        else:
            medication_name, dose = _split_medication_name_dose(_clean_text(record.what))
        quantity = str(record.capsule_quantity) if record.capsule_quantity is not None else ""
        unit = _medication_unit_from_text(medication_name, dose, record.description)
        serialized.append(
            {
                "date": record.date.strftime("%Y-%m-%d"),
                "time": record.time.strftime("%H:%M") if record.time else "",
                "category": record.get_type_display(),
                "created_date": created_at.strftime("%Y-%m-%d"),
                "created_time": created_at.strftime("%H:%M"),
                "recurrence": record.get_recurrence_display() if record.recurrence else "",
                "medication": _clean_text(record.what),
                "medication_name": medication_name,
                "dose": dose,
                "quantity": quantity,
                "quantity_unit": unit,
                "description": _clean_text(record.description),
                "patient": patient_name or str(record.patient),
                "caregiver": _clean_text(record.author_name),
                "status": record.get_status_display(),
                "exception": "Sim" if record.is_exception else "Não",
                "missed_reason": _clean_text(record.missed_reason),
            }
        )
    return serialized


def serialize_bathroom_export(records: Iterable) -> list[dict[str, str]]:
    """Return rows for the bathroom export layout."""
    serialized: list[dict[str, str]] = []
    for record in records:
        serialized.append(
            {
                "date": record.date.strftime("%Y-%m-%d"),
                "time": record.time.strftime("%H:%M") if record.time else "",
                "category": record.get_type_display(),
                "what": _clean_text(record.what_display),
                "description": _clean_text(record.description),
                "caregiver": _clean_text(record.author_name),
                "patient": str(record.patient),
                "status": record.get_status_display(),
                "recurrence": record.get_recurrence_display() if record.recurrence else "",
                "exception": "Sim" if record.is_exception else "Não",
            }
        )
    return serialized


def serialize_vital_export(records: Iterable) -> list[dict[str, str]]:
    """Return rows for the vital signs export layout."""
    serialized: list[dict[str, str]] = []
    for record in records:
        serialized.append(
            {
                "date": record.date.strftime("%Y-%m-%d"),
                "time": record.time.strftime("%H:%M") if record.time else "",
                "category": record.get_type_display(),
                "what": _clean_text(record.what_display),
                "description": _clean_text(record.description),
                "caregiver": _clean_text(record.author_name),
                "patient": str(record.patient),
                "status": record.get_status_display(),
                "recurrence": record.get_recurrence_display() if record.recurrence else "",
                "exception": "Sim" if record.is_exception else "Não",
            }
        )
    return serialized


def serialize_activity_export(records: Iterable) -> list[dict[str, str]]:
    """Return rows for the physical activity export layout."""
    serialized: list[dict[str, str]] = []
    for record in records:
        serialized.append(
            {
                "date": record.date.strftime("%Y-%m-%d"),
                "time": record.time.strftime("%H:%M") if record.time else "",
                "category": record.get_type_display(),
                "what": _clean_text(record.what_display),
                "description": (record.description or "").replace("\r\n", "\n").replace("\r", "\n").strip(),
                "caregiver": _clean_text(record.author_name),
                "patient": str(record.patient),
                "status": record.get_status_display(),
                "recurrence": record.get_recurrence_display() if record.recurrence else "",
                "exception": "Sim" if record.is_exception else "Não",
                "missed_reason": (record.missed_reason or "").replace("\r\n", "\n").replace("\r", "\n").strip(),
            }
        )
    return serialized


def serialize_meal_export(records: Iterable) -> list[dict[str, str]]:
    """Return rows for the meal export layout."""
    serialized: list[dict[str, str]] = []
    for record in records:
        serialized.append(
            {
                "date": record.date.strftime("%Y-%m-%d"),
                "time": record.time.strftime("%H:%M") if record.time else "",
                "category": record.get_type_display(),
                "what": _clean_text(record.what_display),
                "description": (record.description or "").replace("\r\n", "\n").replace("\r", "\n").strip(),
                "caregiver": _clean_text(record.author_name),
                "patient": str(record.patient),
                "status": record.get_status_display(),
                "recurrence": record.get_recurrence_display() if record.recurrence else "",
                "exception": "Sim" if record.is_exception else "Não",
            }
        )
    return serialized


def serialize_progress_export(records: Iterable) -> list[dict[str, str]]:
    """Return rows for the evolution/regression export layout."""
    serialized: list[dict[str, str]] = []
    for record in records:
        serialized.append(
            {
                "date": record.date.strftime("%Y-%m-%d"),
                "time": record.time.strftime("%H:%M") if record.time else "",
                "category": record.get_type_display(),
                "what": _clean_text(record.what_display),
                "description": (record.description or "").replace("\r\n", "\n").replace("\r", "\n").strip(),
                "caregiver": _clean_text(record.author_name),
                "patient": str(record.patient),
                "status": record.get_status_display(),
                "recurrence": record.get_recurrence_display() if record.recurrence else "",
                "exception": "Sim" if record.is_exception else "Não",
                "progress_trend": _clean_text(record.get_progress_trend_display() if record.progress_trend else ""),
            }
        )
    return serialized


def _default_filename(meta: ExportMetadata, ext: str) -> str:
    timestamp = timezone.localtime().strftime("%Y%m%d_%H%M%S")
    return f"registros_{meta.range_slug}_{meta.patient_slug}_{timestamp}.{ext}"


CONSOLIDATED_TYPE_TITLES = {
    "sleep": "Sono",
    "bathroom": "Banheiro e Higiene",
    "vital": "Sinais Vitais",
    "medication": "Medicação",
    "activity": "Atividade Física",
    "meal": "Alimentação",
    "progress": "Evolução / Regressão",
    "other": "Outros",
}

CONSOLIDATED_TYPE_DESCRIPTIONS = {
    "sleep": "Consolidação de períodos de sono, despertar, duração e observações clínicas.",
    "bathroom": "Monitoramento de higiene, eliminações fisiológicas e intercorrências.",
    "vital": "Triagem e acompanhamento de parâmetros clínicos vitais.",
    "medication": "Controle de administração, dose, quantidade e intercorrências medicamentosas.",
    "activity": "Registro de mobilidade, condicionamento e reabilitação motora.",
    "meal": "Acompanhamento de refeições, aceitação alimentar e ocorrências nutricionais.",
    "progress": "Registro de evolução, regressão e mudanças funcionais ou comportamentais.",
    "other": "Registros complementares de cuidado.",
}

CONSOLIDATED_SUMMARY_TITLES = {
    "sleep": "Resumo do Sono",
    "bathroom": "Resumo das Ocorrências",
    "vital": "Resumo de Alertas Clínicos",
    "medication": "Resumo de Ministrações",
    "activity": "Resumo do Período",
    "meal": "Resumo Nutricional",
    "progress": "Resumo de Quadros Observados",
}

CONSOLIDATED_NOTES = {
    "sleep": "Registros consolidados de repouso, despertares e qualidade do sono.",
    "bathroom": "Monitoramento diário de rotinas de higiene fisiológica e cuidados básicos de saúde.",
    "vital": "Triagem e monitoramento preventivo de parâmetros hemodinâmicos e térmicos.",
    "medication": "Controle rigoroso de administração de fármacos, horários e posologias recomendadas.",
    "activity": "Registro e evolução do plano de mobilidade, condicionamento e reabilitação motora.",
    "meal": "Acompanhamento diário do apetite, consistência alimentar e resposta a dietas específicas.",
    "progress": "Consolidado do comportamento, autonomia e parâmetros de recuperação neurológica ou motora.",
}


def _section_meta(meta: ExportMetadata, label: str, total: int) -> ExportMetadata:
    return ExportMetadata(
        start=meta.start,
        end=meta.end,
        period_label=meta.period_label,
        patient_name=meta.patient_name,
        records_total=total,
        group_name=meta.group_name,
        record_types_label=label,
        patient_identifier=meta.patient_identifier,
        professional_name=meta.professional_name,
        unit_name=meta.unit_name,
    )


def _build_consolidated_layout(section: ConsolidatedExportSection, meta: ExportMetadata):
    section_meta = _section_meta(meta, section.label, len(section.rows))
    if section.type_value == "sleep":
        return build_sleep_export_layout(section.rows, section_meta)
    if section.type_value == "bathroom":
        return build_bathroom_export_layout(section.rows, section_meta)
    if section.type_value == "vital":
        return build_vital_export_layout(section.rows, section_meta)
    if section.type_value == "medication":
        return build_medication_export_layout(section.rows, section_meta)
    if section.type_value == "activity":
        return build_activity_export_layout(section.rows, section_meta)
    if section.type_value == "meal":
        return build_meal_export_layout(section.rows, section_meta)
    if section.type_value == "progress":
        return build_progress_export_layout(section.rows, section_meta)
    return None


def _layout_summary_cards(layout) -> list[tuple[str, str, str]]:
    cards = getattr(layout, "summary_cards", [])
    normalized: list[tuple[str, str, str]] = []
    for index, card in enumerate(cards):
        if len(card) == 3:
            label, value, kind = card
        else:
            label, value = card
            kind = "main" if index == 0 else "blue"
        normalized.append((label, value, kind))
    return normalized


def _layout_history_columns(section_type: str, layout) -> list[str]:
    if layout is not None:
        return list(layout.history_columns)
    return [label.upper() for _key, label in COLUMNS]


def _layout_history_rows(section_type: str, layout, rows: list[dict[str, str]]) -> list[dict[str, object]]:
    history: list[dict[str, object]] = []
    if layout is None:
        for row in rows:
            history.append(
                {
                    "values": [row.get(key, "") or "Não informado" for key, _label in COLUMNS],
                    "tag_index": None,
                    "tag_kind": "",
                    "empty": False,
                }
            )
        return history

    for row in layout.history_rows:
        if section_type == "sleep":
            values = [row["date"], row["start"] or "Não informado", row["end"] or "Não informado", row["duration"], row["status"], row["observation"]]
            tag_index = 4
            tag_kind = row.get("status_key", "")
        elif section_type == "bathroom":
            values = [row["date"], row["time"], row["occurrence"], row["recurrence"], row["observation"]]
            tag_index = 2
            tag_kind = row.get("tag_kind", "")
        elif section_type == "vital":
            values = [row["date"], row["time"], row["vital"], row["recurrence"], row["status"], row["observation"]]
            tag_index = 4
            tag_kind = row.get("status_level", "")
        elif section_type == "medication":
            med_detail = row["medication_name"]
            if row.get("dose"):
                med_detail = f"{med_detail}\n{row['dose']}"
            values = [row["date"], row["time"], med_detail, row["quantity"], row["recurrence"], row["observation"]]
            tag_index = 3
            tag_kind = "neutral"
        elif section_type == "activity":
            values = [row["date"], row["time"], row["activity_name"], row["recurrence"], row["observation"]]
            tag_index = None
            tag_kind = ""
        elif section_type == "meal":
            values = [row["date"], row["time"], row["meal_name"], row["recurrence"], row["acceptance"], row["observation"]]
            tag_index = 4
            tag_kind = row.get("acceptance_kind", "")
        elif section_type == "progress":
            values = [row["date"], row["time"], row["classification"], row["recurrence"], row["observation"]]
            tag_index = 2
            tag_kind = row.get("classification_kind", "")
        else:
            values = []
            tag_index = None
            tag_kind = ""
        history.append({"values": values, "tag_index": tag_index, "tag_kind": tag_kind, "empty": False})
    return history


def _empty_section_message(section_title: str, column_count: int) -> dict[str, object]:
    values = [f"Nenhum registro encontrado para {section_title.lower()} no período analisado."]
    values.extend(["" for _ in range(max(column_count - 1, 0))])
    return {"values": values, "tag_index": None, "tag_kind": "", "empty": True}


def _summary_kind_color(kind: str) -> str:
    return {
        "normal": "22543D",
        "good": "22543D",
        "evolution": "22543D",
        "warning": "B7791F",
        "alert": "9B2C2C",
        "poor": "9B2C2C",
        "regression": "9B2C2C",
        "neutral": "4A5568",
        "main": "2B6CB0",
    }.get(kind, "2B6CB0")


def _tag_colors(kind: str) -> tuple[str, str]:
    return {
        "done": ("C6F6D5", "22543D"),
        "pending": ("FEFCBF", "744210"),
        "missed": ("FED7D7", "9B2C2C"),
        "hygiene": ("DCEEFF", "1D5F91"),
        "elimination": ("E5E7EB", "374151"),
        "alert": ("FED7D7", "9B2C2C"),
        "normal": ("C6F6D5", "22543D"),
        "warning": ("FEFCBF", "744210"),
        "good": ("C6F6D5", "22543D"),
        "poor": ("FED7D7", "9B2C2C"),
        "evolution": ("C6F6D5", "22543D"),
        "regression": ("FED7D7", "9B2C2C"),
        "neutral": ("E2E8F0", "4A5568"),
        "other": ("E2E8F0", "4A5568"),
    }.get(kind, ("E2E8F0", "4A5568"))


def export_as_csv(
    rows: list[dict[str, str]],
    meta: ExportMetadata,
    columns: Sequence[tuple[str, str]] = COLUMNS,
) -> HttpResponse:
    response = HttpResponse(content_type="text/csv; charset=utf-8")
    response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'csv')}\""
    response.write("\ufeff")
    writer = csv.writer(response)
    writer.writerow([label for _, label in columns])
    for row in rows:
        writer.writerow([row[key] for key, _ in columns])
    return response


def export_as_xlsx(
    rows: list[dict[str, str]],
    meta: ExportMetadata,
    columns: Sequence[tuple[str, str]] = COLUMNS,
) -> HttpResponse:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        return _export_xlsx_inline(rows, meta)

    wb = Workbook()
    ws = wb.active
    ws.title = "Registros"

    total_columns = len(columns)
    header_row = 1
    header_font = Font(bold=True, color="000000")
    thin_side = Side(style="thin", color="D1D5DB")
    header_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    cell_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    for col_idx, (_, label) in enumerate(columns, start=1):
        cell = ws.cell(row=header_row, column=col_idx, value=label)
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = header_border

    column_widths = _column_width_units(columns)
    for data_idx, row in enumerate(rows, start=1):
        excel_row = header_row + data_idx
        for col_idx, (key, _) in enumerate(columns, start=1):
            value = _xlsx_sanitize(row.get(key, ""))
            cell = ws.cell(row=excel_row, column=col_idx, value=value)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = cell_border

    ws.freeze_panes = ws.cell(row=2, column=1)
    ws.row_dimensions[1].height = 22
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width

    buffer = BytesIO()
    wb.save(buffer)
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'xlsx')}\""
    return response


def export_consolidated_as_docx(
    sections: list[ConsolidatedExportSection],
    meta: ExportMetadata,
) -> HttpResponse:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise ExportDependencyError("Exportação DOCX indisponível.") from exc

    def set_cell_shading(cell, color: str) -> None:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_border(cell, color: str = "D9E2EC", size: str = "8") -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            tag = f"w:{edge}"
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), size)
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color)

    def set_paragraph_text(paragraph, text: str, *, bold: bool = False, color: str | None = None, size: int | None = None) -> None:
        run = paragraph.add_run(text)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if size:
            run.font.size = Pt(size)

    def add_rule(document) -> None:
        line = document.add_table(rows=1, cols=1)
        line.style = "Table Grid"
        set_cell_shading(line.rows[0].cells[0], "E2E8F0")
        set_cell_border(line.rows[0].cells[0], "E2E8F0", "4")

    def add_section_heading(document, text: str) -> None:
        paragraph = document.add_paragraph()
        set_paragraph_text(paragraph, text, bold=True, color="2B6CB0", size=14)
        add_rule(document)

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    header_table = document.add_table(rows=1, cols=2)
    header_table.style = "Table Grid"
    for cell in header_table.rows[0].cells:
        set_cell_border(cell, "FFFFFF", "0")
    left, right = header_table.rows[0].cells
    set_paragraph_text(left.paragraphs[0], "Relatório Consolidado de Cuidados", bold=True, color="1A365D", size=22)
    subtitle = left.add_paragraph()
    set_paragraph_text(subtitle, "Cuidar Juntos", color="4A5568", size=10)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_text(right.paragraphs[0], f"Geração: {timezone.localtime().strftime('%d/%m/%Y %H:%M')}", color="718096", size=9)
    add_rule(document)

    document.add_paragraph()
    info_table = document.add_table(rows=1, cols=4)
    info_table.style = "Table Grid"
    info_items = [
        ("PACIENTE", meta.patient_name or "Todos os pacientes"),
        ("PROFISSIONAL RESPONSÁVEL", meta.professional_name or "Não informado"),
        ("PERÍODO ANALISADO", meta.period_label or "Todos os tempos"),
        ("TIPOS DE REGISTRO EXPORTADOS", meta.record_types_label or "Todos os tipos"),
    ]
    for idx, (label, value) in enumerate(info_items):
        cell = info_table.rows[0].cells[idx]
        set_cell_border(cell, "EDF2F7", "8")
        set_cell_shading(cell, "F7FAFC")
        set_paragraph_text(cell.paragraphs[0], label, bold=True, color="718096", size=7)
        paragraph = cell.add_paragraph()
        set_paragraph_text(paragraph, value, bold=True, color="2D3748", size=9)

    add_section_heading(document, "Resumo Geral da Exportação")
    summary_table = document.add_table(rows=1, cols=4)
    summary_table.style = "Table Grid"
    summary_items = [
        ("Tipos Exportados", str(len(sections)), "main"),
        ("Registros Avaliados", str(meta.records_total), "blue"),
        ("Tipos Selecionados", meta.record_types_label or "Todos os tipos", "blue"),
        ("Período Analisado", meta.period_label or "Todos os tempos", "neutral"),
    ]
    for idx, (label, value, kind) in enumerate(summary_items):
        cell = summary_table.rows[0].cells[idx]
        set_cell_border(cell, "BEE3F8" if kind == "main" else "E2E8F0", "8")
        set_cell_shading(cell, "EBF8FF" if kind == "main" else "FFFFFF")
        set_paragraph_text(cell.paragraphs[0], label, bold=True, color="4A5568", size=7)
        paragraph = cell.add_paragraph()
        set_paragraph_text(paragraph, value, bold=True, color=_summary_kind_color(kind), size=14 if idx in (2, 3) else 16)

    for section_index, section_data in enumerate(sections):
        if section_index > 0:
            document.add_page_break()
        section_title = CONSOLIDATED_TYPE_TITLES.get(section_data.type_value, section_data.label)
        title_paragraph = document.add_paragraph()
        set_paragraph_text(title_paragraph, section_title, bold=True, color="1A365D", size=18)
        desc = document.add_paragraph()
        set_paragraph_text(desc, CONSOLIDATED_TYPE_DESCRIPTIONS.get(section_data.type_value, ""), color="4A5568", size=9)
        add_rule(document)

        layout = _build_consolidated_layout(section_data, meta)
        add_section_heading(document, CONSOLIDATED_SUMMARY_TITLES.get(section_data.type_value, "Resumo do Tipo"))
        cards = _layout_summary_cards(layout) if layout is not None else [("Total de Registros", str(len(section_data.rows)), "main")]
        card_table = document.add_table(rows=1, cols=max(len(cards), 1))
        card_table.style = "Table Grid"
        for idx, (label, value, kind) in enumerate(cards):
            cell = card_table.rows[0].cells[idx]
            set_cell_border(cell, "BEE3F8" if kind == "main" else "E2E8F0", "8")
            set_cell_shading(cell, "EBF8FF" if kind == "main" else "FFFFFF")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="4A5568", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value, bold=True, color=_summary_kind_color(kind), size=16 if kind == "main" else 15)
        if section_data.type_value in CONSOLIDATED_NOTES:
            note = document.add_paragraph()
            set_paragraph_text(note, CONSOLIDATED_NOTES[section_data.type_value], color="718096", size=9)

        add_section_heading(document, "Histórico Detalhado")
        columns = _layout_history_columns(section_data.type_value, layout)
        history_rows = _layout_history_rows(section_data.type_value, layout, section_data.rows)
        if not history_rows:
            history_rows.append(_empty_section_message(section_title, len(columns)))
        table = document.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        for idx, label in enumerate(columns):
            cell = table.rows[0].cells[idx]
            set_cell_shading(cell, "2B6CB0")
            set_cell_border(cell, "2B6CB0", "6")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="FFFFFF", size=7)
        for row_index, row in enumerate(history_rows, start=1):
            cells = table.add_row().cells
            values = row["values"]
            tag_index = row["tag_index"]
            tag_kind = row["tag_kind"]
            for idx, value in enumerate(values):
                set_cell_border(cells[idx], "E2E8F0", "4")
                if row_index % 2 == 0:
                    set_cell_shading(cells[idx], "F7FAFC")
                paragraph = cells[idx].paragraphs[0]
                if idx == 0:
                    set_paragraph_text(paragraph, str(value), bold=True, color="2D3748", size=8)
                elif tag_index is not None and idx == tag_index and not row.get("empty"):
                    fill, color = _tag_colors(str(tag_kind))
                    set_cell_shading(cells[idx], fill)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_text(paragraph, str(value).upper(), bold=True, color=color, size=7)
                else:
                    set_paragraph_text(paragraph, str(value or ""), color="4A5568" if idx == len(values) - 1 else "2D3748", size=7 if idx == len(values) - 1 else 8)

        technical_notes = getattr(layout, "technical_notes", []) if layout is not None else []
        if technical_notes:
            add_section_heading(document, "Informações Técnicas")
            technical_table = document.add_table(rows=1, cols=2)
            technical_table.style = "Table Grid"
            bar, content = technical_table.rows[0].cells
            set_cell_shading(bar, "CBD5E0")
            set_cell_border(bar, "CBD5E0", "4")
            set_cell_shading(content, "F7FAFC")
            set_cell_border(content, "F7FAFC", "4")
            for note_text in technical_notes:
                paragraph = content.add_paragraph()
                set_paragraph_text(paragraph, note_text, color="4A5568", size=8)

    buffer = BytesIO()
    document.save(buffer)
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'docx')}\""
    return response


def export_consolidated_as_pdf(
    sections: list[ConsolidatedExportSection],
    meta: ExportMetadata,
) -> HttpResponse:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import HRFlowable, PageBreak, Paragraph, SimpleDocTemplate, Table, TableStyle
    except ImportError as exc:
        raise ExportDependencyError("Exportação PDF indisponível.") from exc

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, title="Relatório Consolidado de Cuidados")
    base_styles = getSampleStyleSheet()
    story: list[object] = []
    blue_dark = colors.HexColor("#1A365D")
    blue = colors.HexColor("#2B6CB0")
    gray_line = colors.HexColor("#E2E8F0")
    gray_text = colors.HexColor("#4A5568")
    gray_muted = colors.HexColor("#718096")
    gray_soft = colors.HexColor("#F7FAFC")

    title_style = ParagraphStyle("ConsolidatedTitle", parent=base_styles["Title"], fontName="Helvetica-Bold", fontSize=22, leading=27, textColor=blue_dark, alignment=0)
    subtitle_style = ParagraphStyle("ConsolidatedSubtitle", parent=base_styles["Normal"], fontName="Helvetica", fontSize=10, leading=12, textColor=gray_text)
    generation_style = ParagraphStyle("ConsolidatedGeneration", parent=base_styles["Normal"], fontName="Helvetica", fontSize=8, leading=10, alignment=2, textColor=gray_muted)
    section_title_style = ParagraphStyle("ConsolidatedSectionTitle", parent=base_styles["Heading2"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=blue_dark, spaceBefore=10, spaceAfter=3)
    subsection_style = ParagraphStyle("ConsolidatedSubsection", parent=base_styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=blue, spaceBefore=16, spaceAfter=5)
    label_style = ParagraphStyle("ConsolidatedLabel", parent=base_styles["Normal"], fontName="Helvetica-Bold", fontSize=7, leading=9, textColor=gray_muted, alignment=1)
    value_style = ParagraphStyle("ConsolidatedValue", parent=base_styles["Normal"], fontName="Helvetica-Bold", fontSize=9, leading=12, textColor=colors.HexColor("#2D3748"), alignment=1, wordWrap="CJK")
    small_style = ParagraphStyle("ConsolidatedSmall", parent=base_styles["Normal"], fontName="Helvetica", fontSize=8, leading=10, textColor=gray_text, wordWrap="CJK")
    table_header_style = ParagraphStyle("ConsolidatedTableHeader", parent=base_styles["Normal"], fontName="Helvetica-Bold", fontSize=7, leading=9, textColor=colors.white, alignment=1)
    table_cell_style = ParagraphStyle("ConsolidatedTableCell", parent=base_styles["Normal"], fontName="Helvetica", fontSize=7, leading=9, textColor=colors.HexColor("#2D3748"), wordWrap="CJK")
    table_bold_style = ParagraphStyle("ConsolidatedTableBold", parent=table_cell_style, fontName="Helvetica-Bold")

    header_table = Table(
        [[[Paragraph("Relatório Consolidado de Cuidados", title_style), Paragraph("Cuidar Juntos", subtitle_style)], Paragraph(f"Geração: {timezone.localtime().strftime('%d/%m/%Y %H:%M')}", generation_style)]],
        colWidths=[doc.width * 0.66, doc.width * 0.34],
    )
    header_table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0)]))
    story.append(header_table)
    story.append(HRFlowable(width="100%", thickness=1.0, color=gray_line, spaceBefore=2, spaceAfter=18))

    info_items = [
        ("PACIENTE", meta.patient_name or "Todos os pacientes"),
        ("PROFISSIONAL RESPONSÁVEL", meta.professional_name or "Não informado"),
        ("PERÍODO ANALISADO", meta.period_label or "Todos os tempos"),
        ("TIPOS DE REGISTRO EXPORTADOS", meta.record_types_label or "Todos os tipos"),
    ]
    info_table = Table([[[Paragraph(label, label_style), Paragraph(value, value_style)] for label, value in info_items]], colWidths=[doc.width / 4] * 4)
    info_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), gray_soft), ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#EDF2F7")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 10), ("BOTTOMPADDING", (0, 0), (-1, -1), 10)]))
    story.append(info_table)

    story.append(Paragraph("Resumo Geral da Exportação", subsection_style))
    story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
    summary_items = [
        ("Tipos Exportados", str(len(sections)), "main"),
        ("Registros Avaliados", str(meta.records_total), "blue"),
        ("Tipos Selecionados", meta.record_types_label or "Todos os tipos", "blue"),
        ("Período Analisado", meta.period_label or "Todos os tempos", "neutral"),
    ]
    summary_cells = []
    for label, value, kind in summary_items:
        card_value_style = ParagraphStyle(f"ConsolidatedSummary{kind}", parent=value_style, fontSize=14, leading=17, textColor=colors.HexColor(f"#{_summary_kind_color(kind)}"))
        summary_cells.append([Paragraph(label, label_style), Paragraph(value, card_value_style)])
    summary_table = Table([summary_cells], colWidths=[doc.width / 4] * 4)
    summary_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#EBF8FF")), ("BOX", (0, 0), (0, 0), 0.8, colors.HexColor("#BEE3F8")), ("BOX", (1, 0), (-1, -1), 0.7, gray_line), ("INNERGRID", (0, 0), (-1, -1), 0.5, gray_line), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 10), ("BOTTOMPADDING", (0, 0), (-1, -1), 10)]))
    story.append(summary_table)

    for section_index, section_data in enumerate(sections):
        if section_index > 0:
            story.append(PageBreak())
        section_title = CONSOLIDATED_TYPE_TITLES.get(section_data.type_value, section_data.label)
        story.append(Paragraph(section_title, section_title_style))
        description = CONSOLIDATED_TYPE_DESCRIPTIONS.get(section_data.type_value, "")
        if description:
            story.append(Paragraph(description, small_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=2, spaceAfter=10))

        layout = _build_consolidated_layout(section_data, meta)
        story.append(Paragraph(CONSOLIDATED_SUMMARY_TITLES.get(section_data.type_value, "Resumo do Tipo"), subsection_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        cards = _layout_summary_cards(layout) if layout is not None else [("Total de Registros", str(len(section_data.rows)), "main")]
        card_cells = []
        for label, value, kind in cards:
            card_value_style = ParagraphStyle(f"ConsolidatedCard{section_data.type_value}{kind}", parent=value_style, fontSize=15, leading=18, textColor=colors.HexColor(f"#{_summary_kind_color(kind)}"))
            card_cells.append([Paragraph(label, label_style), Paragraph(value, card_value_style)])
        card_table = Table([card_cells], colWidths=[doc.width / max(len(cards), 1)] * max(len(cards), 1))
        card_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#EBF8FF")), ("BOX", (0, 0), (0, 0), 0.8, colors.HexColor("#BEE3F8")), ("BOX", (1, 0), (-1, -1), 0.7, gray_line), ("INNERGRID", (0, 0), (-1, -1), 0.5, gray_line), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 10), ("BOTTOMPADDING", (0, 0), (-1, -1), 10)]))
        story.append(card_table)
        if section_data.type_value in CONSOLIDATED_NOTES:
            story.append(Paragraph(CONSOLIDATED_NOTES[section_data.type_value], small_style))

        story.append(Paragraph("Histórico Detalhado", subsection_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        columns = _layout_history_columns(section_data.type_value, layout)
        history_rows = _layout_history_rows(section_data.type_value, layout, section_data.rows)
        if not history_rows:
            history_rows.append(_empty_section_message(section_title, len(columns)))
        data: list[list[object]] = [[Paragraph(label, table_header_style) for label in columns]]
        for row in history_rows:
            row_values = []
            for idx, value in enumerate(row["values"]):
                style = table_bold_style if idx == 0 else table_cell_style
                row_values.append(Paragraph(xml_escape(str(value or "")).replace("\n", "<br/>"), style))
            data.append(row_values)
        unit_total = len(columns) or 1
        table = Table(data, repeatRows=1, colWidths=[doc.width / unit_total] * unit_total)
        style_commands = [
            ("BACKGROUND", (0, 0), (-1, 0), blue),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, 0), (-1, -1), 0.5, gray_line),
        ]
        for idx, row in enumerate(history_rows, start=1):
            if idx % 2 == 0:
                style_commands.append(("BACKGROUND", (0, idx), (-1, idx), gray_soft))
            tag_index = row["tag_index"]
            if tag_index is not None and not row.get("empty"):
                bg, _fg = _tag_colors(str(row["tag_kind"]))
                style_commands.append(("BACKGROUND", (tag_index, idx), (tag_index, idx), colors.HexColor(f"#{bg}")))
                style_commands.append(("ALIGN", (tag_index, idx), (tag_index, idx), "CENTER"))
        table.setStyle(TableStyle(style_commands))
        story.append(table)

        technical_notes = getattr(layout, "technical_notes", []) if layout is not None else []
        if technical_notes:
            story.append(Paragraph("Informações Técnicas", subsection_style))
            story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
            technical_content = [[Paragraph(note, small_style)] for note in technical_notes]
            technical_table = Table([[Table([[""]], colWidths=[4]), Table(technical_content, colWidths=[doc.width - 28])]], colWidths=[8, doc.width - 8])
            technical_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#CBD5E0")), ("BACKGROUND", (1, 0), (1, 0), gray_soft), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
            story.append(technical_table)

    doc.build(story)
    response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'pdf')}\""
    return response


def export_as_docx(
    rows: list[dict[str, str]],
    meta: ExportMetadata,
    columns: Sequence[tuple[str, str]] = COLUMNS,
) -> HttpResponse:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        return _export_docx_inline(rows, meta, columns=columns)

    def set_cell_shading(cell, color: str) -> None:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_border(cell, color: str = "D9E2EC", size: str = "8") -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            tag = f"w:{edge}"
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), size)
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color)

    def set_paragraph_text(paragraph, text: str, *, bold: bool = False, italic: bool = False, color: str | None = None, size: int | None = None) -> None:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if size:
            run.font.size = Pt(size)

    if is_sleep_export(rows, meta):
        layout = build_sleep_export_layout(rows, meta)
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

        header_table = document.add_table(rows=1, cols=2)
        header_table.autofit = True
        header_table.style = "Table Grid"
        for cell in header_table.rows[0].cells:
            set_cell_border(cell, "FFFFFF", "0")
        left, right = header_table.rows[0].cells
        set_paragraph_text(left.paragraphs[0], layout.title, bold=True, color="123A5A", size=22)
        sub = left.add_paragraph()
        set_paragraph_text(sub, "Cuidar Juntos", bold=True, color="4B5563", size=10)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_text(right.paragraphs[0], f"Geração: {layout.generated_at}", color="4B5563", size=9)

        divider = document.add_table(rows=1, cols=1)
        divider.style = "Table Grid"
        set_cell_shading(divider.rows[0].cells[0], "D9E2EC")
        set_cell_border(divider.rows[0].cells[0], "D9E2EC", "4")

        document.add_paragraph()
        info_table = document.add_table(rows=1, cols=3)
        info_table.style = "Table Grid"
        for idx, (label, value) in enumerate(
            [
                ("PACIENTE", layout.patient_label),
                ("PROFISSIONAL RESPONSÁVEL", layout.professional_label),
                ("PERÍODO ANALISADO", layout.period_label),
            ]
        ):
            cell = info_table.rows[0].cells[idx]
            set_cell_border(cell, "D9E2EC", "8")
            set_cell_shading(cell, "FFFFFF")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="6B7280", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value or "Não informado", bold=True, color="111827", size=10)

        def add_section_title(text: str) -> None:
            paragraph = document.add_paragraph()
            set_paragraph_text(paragraph, text, bold=True, color="1D5F91", size=14)
            line = document.add_table(rows=1, cols=1)
            line.style = "Table Grid"
            set_cell_shading(line.rows[0].cells[0], "D9E2EC")
            set_cell_border(line.rows[0].cells[0], "D9E2EC", "4")

        add_section_title("Resumo do Sono")
        summary_table = document.add_table(rows=1, cols=4)
        summary_table.style = "Table Grid"
        for idx, (label, value) in enumerate(layout.summary_cards):
            cell = summary_table.rows[0].cells[idx]
            set_cell_border(cell, "B7D7F2" if idx == 0 else "D9E2EC", "10")
            set_cell_shading(cell, "F8FBFF" if idx == 0 else "FFFFFF")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="6B7280", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value, bold=True, color="1D5F91", size=16)

        summary_line = document.add_paragraph()
        set_paragraph_text(
            summary_line,
            (
                f"Registros de início do sono: {layout.start_count} | "
                f"Registros de término/despertar: {layout.end_count} | "
                f"Total de registros avaliados: {layout.total_count}"
            ),
            color="4B5563",
            size=9,
        )

        add_section_title("Histórico Detalhado")
        history_table = document.add_table(rows=1, cols=len(layout.history_columns))
        history_table.style = "Table Grid"
        for idx, label in enumerate(layout.history_columns):
            cell = history_table.rows[0].cells[idx]
            set_cell_shading(cell, "F3F4F6")
            set_cell_border(cell, "D9E2EC", "6")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="6B7280", size=7)

        for row in layout.history_rows:
            cells = history_table.add_row().cells
            values = [
                row["date"],
                row["start"],
                row["end"],
                row["duration"],
                row["status"],
                row["observation"],
            ]
            for idx, value in enumerate(values):
                set_cell_border(cells[idx], "E5E7EB", "4")
                paragraph = cells[idx].paragraphs[0]
                if idx == 0:
                    set_paragraph_text(paragraph, value or "Não informado", bold=True, color="111827", size=8)
                elif idx in (1, 2) and not value:
                    set_paragraph_text(paragraph, "Não informado", italic=True, color="6B7280", size=8)
                elif idx == 3 and not row["duration_complete"]:
                    set_paragraph_text(paragraph, "Registro incompleto", italic=True, color="B91C1C", size=8)
                elif idx == 3:
                    set_paragraph_text(paragraph, value, bold=True, color="111827", size=8)
                elif idx == 4:
                    status_color = {"done": "047857", "pending": "B45309", "missed": "B91C1C"}.get(row["status_key"], "374151")
                    set_paragraph_text(paragraph, value or "Não informado", bold=True, color=status_color, size=8)
                    if row["caregiver"]:
                        resp = cells[idx].add_paragraph()
                        set_paragraph_text(resp, f"Resp: {row['caregiver']}", color="6B7280", size=7)
                else:
                    set_paragraph_text(paragraph, value or "", color="374151", size=8)

        add_section_title("Informações Técnicas")
        technical_table = document.add_table(rows=1, cols=2)
        technical_table.style = "Table Grid"
        bar, content = technical_table.rows[0].cells
        set_cell_shading(bar, "B7D7F2")
        set_cell_border(bar, "B7D7F2", "4")
        set_cell_border(content, "E5E7EB", "4")
        for note in layout.technical_notes:
            paragraph = content.add_paragraph()
            set_paragraph_text(paragraph, note, color="374151", size=8)

        buffer = BytesIO()
        document.save(buffer)
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'docx')}\""
        return response

    if is_bathroom_export(rows, meta):
        layout = build_bathroom_export_layout(rows, meta)
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

        header_table = document.add_table(rows=1, cols=2)
        header_table.style = "Table Grid"
        for cell in header_table.rows[0].cells:
            set_cell_border(cell, "FFFFFF", "0")
        left, right = header_table.rows[0].cells
        set_paragraph_text(left.paragraphs[0], layout.title, bold=True, color="123A5A", size=21)
        sub = left.add_paragraph()
        set_paragraph_text(sub, "Cuidar Juntos", bold=True, color="4B5563", size=10)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_text(right.paragraphs[0], f"Geração: {layout.generated_at}", color="4B5563", size=9)

        divider = document.add_table(rows=1, cols=1)
        divider.style = "Table Grid"
        set_cell_shading(divider.rows[0].cells[0], "D9E2EC")
        set_cell_border(divider.rows[0].cells[0], "D9E2EC", "4")

        document.add_paragraph()
        info_table = document.add_table(rows=1, cols=3)
        info_table.style = "Table Grid"
        for idx, (label, value) in enumerate(
            [
                ("PACIENTE", layout.patient_label),
                ("PROFISSIONAL RESPONSÁVEL", layout.professional_label),
                ("PERÍODO ANALISADO", layout.period_label),
            ]
        ):
            cell = info_table.rows[0].cells[idx]
            set_cell_border(cell, "D9E2EC", "8")
            set_cell_shading(cell, "F8FAFC")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="6B7280", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value or "Não informado", bold=True, color="111827", size=10)

        def add_bathroom_section(text: str) -> None:
            paragraph = document.add_paragraph()
            set_paragraph_text(paragraph, text, bold=True, color="1D5F91", size=14)
            line = document.add_table(rows=1, cols=1)
            line.style = "Table Grid"
            set_cell_shading(line.rows[0].cells[0], "D9E2EC")
            set_cell_border(line.rows[0].cells[0], "D9E2EC", "4")

        add_bathroom_section("Resumo das Ocorrências")
        summary_table = document.add_table(rows=1, cols=4)
        summary_table.style = "Table Grid"
        for idx, (label, value) in enumerate(layout.summary_cards):
            cell = summary_table.rows[0].cells[idx]
            set_cell_border(cell, "D9E2EC", "8")
            set_cell_shading(cell, "EAF4FF" if idx == 0 else "FFFFFF")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="6B7280", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value, bold=True, color="1D5F91", size=16)
        note = document.add_paragraph()
        set_paragraph_text(note, "Monitoramento diário de rotinas de higiene fisiológica e cuidados básicos de saúde.", color="4B5563", size=9)

        add_bathroom_section("Histórico Detalhado")
        history_table = document.add_table(rows=1, cols=len(layout.history_columns))
        history_table.style = "Table Grid"
        for idx, label in enumerate(layout.history_columns):
            cell = history_table.rows[0].cells[idx]
            set_cell_shading(cell, "1D5F91")
            set_cell_border(cell, "1D5F91", "6")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="FFFFFF", size=7)

        tag_styles = {
            "hygiene": ("DCEEFF", "1D5F91"),
            "elimination": ("E5E7EB", "374151"),
            "alert": ("FEE2E2", "B91C1C"),
            "neutral": ("F3F4F6", "4B5563"),
        }
        for row_index, row in enumerate(layout.history_rows, start=1):
            cells = history_table.add_row().cells
            values = [row["date"], row["time"], row["occurrence"], row["recurrence"], row["observation"]]
            for idx, value in enumerate(values):
                set_cell_border(cells[idx], "E5E7EB", "4")
                if row_index % 2 == 0:
                    set_cell_shading(cells[idx], "F8FAFC")
                paragraph = cells[idx].paragraphs[0]
                if idx == 0:
                    set_paragraph_text(paragraph, value, bold=True, color="111827", size=8)
                elif idx == 2:
                    fill, color = tag_styles.get(row["tag_kind"], tag_styles["neutral"])
                    set_cell_shading(cells[idx], fill)
                    set_paragraph_text(paragraph, value.upper(), bold=True, color=color, size=7)
                elif idx == 4:
                    set_paragraph_text(paragraph, value or "", color="374151", size=7)
                else:
                    set_paragraph_text(paragraph, value or "Não informado", color="374151", size=8)

        add_bathroom_section("Informações Técnicas")
        technical_table = document.add_table(rows=1, cols=2)
        technical_table.style = "Table Grid"
        bar, content = technical_table.rows[0].cells
        set_cell_shading(bar, "D1D5DB")
        set_cell_border(bar, "D1D5DB", "4")
        set_cell_shading(content, "F3F4F6")
        set_cell_border(content, "F3F4F6", "4")
        for note_text in layout.technical_notes:
            paragraph = content.add_paragraph()
            set_paragraph_text(paragraph, note_text, color="374151", size=8)

        buffer = BytesIO()
        document.save(buffer)
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'docx')}\""
        return response

    if is_vital_export(rows, meta):
        layout = build_vital_export_layout(rows, meta)
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

        header_table = document.add_table(rows=1, cols=2)
        header_table.style = "Table Grid"
        for cell in header_table.rows[0].cells:
            set_cell_border(cell, "FFFFFF", "0")
        left, right = header_table.rows[0].cells
        set_paragraph_text(left.paragraphs[0], layout.title, bold=True, color="1A365D", size=21)
        subtitle = left.add_paragraph()
        set_paragraph_text(subtitle, "Cuidar Juntos", bold=True, color="4A5568", size=10)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_text(right.paragraphs[0], f"Geração: {layout.generated_at}", color="718096", size=9)

        divider = document.add_table(rows=1, cols=1)
        divider.style = "Table Grid"
        set_cell_shading(divider.rows[0].cells[0], "E2E8F0")
        set_cell_border(divider.rows[0].cells[0], "E2E8F0", "4")

        document.add_paragraph()
        info_table = document.add_table(rows=1, cols=3)
        info_table.style = "Table Grid"
        for idx, (label, value) in enumerate(
            [
                ("PACIENTE", layout.patient_label),
                ("PROFISSIONAL RESPONSÁVEL", layout.professional_label),
                ("PERÍODO ANALISADO", layout.period_label),
            ]
        ):
            cell = info_table.rows[0].cells[idx]
            set_cell_border(cell, "EDF2F7", "8")
            set_cell_shading(cell, "F7FAFC")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="718096", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value or "Não informado", bold=True, color="2D3748", size=10)

        def add_vital_section(text: str) -> None:
            paragraph = document.add_paragraph()
            set_paragraph_text(paragraph, text, bold=True, color="2B6CB0", size=14)
            line = document.add_table(rows=1, cols=1)
            line.style = "Table Grid"
            set_cell_shading(line.rows[0].cells[0], "E2E8F0")
            set_cell_border(line.rows[0].cells[0], "E2E8F0", "4")

        add_vital_section("Resumo de Alertas Clínicos")
        summary_table = document.add_table(rows=1, cols=4)
        summary_table.style = "Table Grid"
        card_colors = {
            "main": ("EBF8FF", "2B6CB0"),
            "normal": ("FFFFFF", "22543D"),
            "warning": ("FFFFFF", "B7791F"),
            "alert": ("FFFFFF", "9B2C2C"),
        }
        for idx, (label, value, kind) in enumerate(layout.summary_cards):
            cell = summary_table.rows[0].cells[idx]
            bg, fg = card_colors[kind]
            set_cell_border(cell, "BEE3F8" if kind == "main" else "E2E8F0", "8")
            set_cell_shading(cell, bg)
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="4A5568", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value, bold=True, color=fg, size=17 if kind == "main" else 16)
        note = document.add_paragraph()
        set_paragraph_text(note, "Triagem e monitoramento preventivo de parâmetros hemodinâmicos e térmicos.", color="718096", size=9)

        add_vital_section("Histórico Detalhado")
        history_table = document.add_table(rows=1, cols=len(layout.history_columns))
        history_table.style = "Table Grid"
        for idx, label in enumerate(layout.history_columns):
            cell = history_table.rows[0].cells[idx]
            set_cell_shading(cell, "2B6CB0")
            set_cell_border(cell, "2B6CB0", "6")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="FFFFFF", size=7)

        status_styles = {
            "normal": ("C6F6D5", "22543D"),
            "warning": ("FEFCBF", "744210"),
            "alert": ("FED7D7", "9B2C2C"),
            "neutral": ("E2E8F0", "4A5568"),
        }
        for row_index, row in enumerate(layout.history_rows, start=1):
            cells = history_table.add_row().cells
            values = [row["date"], row["time"], row["vital"], row["recurrence"], row["status"], row["observation"]]
            for idx, value in enumerate(values):
                set_cell_border(cells[idx], "E2E8F0", "4")
                if row_index % 2 == 0:
                    set_cell_shading(cells[idx], "F7FAFC")
                paragraph = cells[idx].paragraphs[0]
                if idx == 0:
                    set_paragraph_text(paragraph, value, bold=True, color="2D3748", size=8)
                elif idx == 2:
                    set_paragraph_text(paragraph, value, bold=True, color="1A365D", size=8)
                elif idx == 4:
                    fill, color = status_styles.get(row["status_level"], status_styles["neutral"])
                    set_cell_shading(cells[idx], fill)
                    set_paragraph_text(paragraph, value.upper(), bold=True, color=color, size=7)
                elif idx == 5:
                    set_paragraph_text(paragraph, value or "", color="4A5568", size=7)
                else:
                    set_paragraph_text(paragraph, value or "Não informado", color="2D3748", size=8)

        add_vital_section("Informações Técnicas")
        technical_table = document.add_table(rows=1, cols=2)
        technical_table.style = "Table Grid"
        bar, content = technical_table.rows[0].cells
        set_cell_shading(bar, "CBD5E0")
        set_cell_border(bar, "CBD5E0", "4")
        set_cell_shading(content, "F7FAFC")
        set_cell_border(content, "F7FAFC", "4")
        for note_text in layout.technical_notes:
            paragraph = content.add_paragraph()
            set_paragraph_text(paragraph, note_text, color="4A5568", size=8)

        buffer = BytesIO()
        document.save(buffer)
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'docx')}\""
        return response

    if is_medication_export(rows, meta):
        layout = build_medication_export_layout(rows, meta)
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

        header_table = document.add_table(rows=1, cols=2)
        header_table.style = "Table Grid"
        for cell in header_table.rows[0].cells:
            set_cell_border(cell, "FFFFFF", "0")
        left, right = header_table.rows[0].cells
        set_paragraph_text(left.paragraphs[0], layout.title, bold=True, color="1A365D", size=21)
        subtitle = left.add_paragraph()
        set_paragraph_text(subtitle, "Cuidar Juntos", bold=True, color="4A5568", size=10)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_text(right.paragraphs[0], f"Geração: {layout.generated_at}", color="718096", size=9)

        divider = document.add_table(rows=1, cols=1)
        divider.style = "Table Grid"
        set_cell_shading(divider.rows[0].cells[0], "E2E8F0")
        set_cell_border(divider.rows[0].cells[0], "E2E8F0", "4")

        document.add_paragraph()
        info_table = document.add_table(rows=1, cols=3)
        info_table.style = "Table Grid"
        for idx, (label, value) in enumerate(
            [
                ("PACIENTE", layout.patient_label),
                ("PROFISSIONAL RESPONSÁVEL", layout.professional_label),
                ("PERÍODO ANALISADO", layout.period_label),
            ]
        ):
            cell = info_table.rows[0].cells[idx]
            set_cell_border(cell, "EDF2F7", "8")
            set_cell_shading(cell, "F7FAFC")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="718096", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value or "Não informado", bold=True, color="2D3748", size=10)

        def add_medication_section(text: str) -> None:
            paragraph = document.add_paragraph()
            set_paragraph_text(paragraph, text, bold=True, color="2B6CB0", size=14)
            line = document.add_table(rows=1, cols=1)
            line.style = "Table Grid"
            set_cell_shading(line.rows[0].cells[0], "E2E8F0")
            set_cell_border(line.rows[0].cells[0], "E2E8F0", "4")

        add_medication_section("Resumo de Ministrações")
        summary_table = document.add_table(rows=1, cols=4)
        summary_table.style = "Table Grid"
        for idx, (label, value, kind) in enumerate(layout.summary_cards):
            cell = summary_table.rows[0].cells[idx]
            set_cell_border(cell, "BEE3F8" if kind == "main" else "E2E8F0", "8")
            set_cell_shading(cell, "EBF8FF" if kind == "main" else "FFFFFF")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="4A5568", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value, bold=True, color="4A5568" if kind == "neutral" else "2B6CB0", size=17 if kind == "main" else 16)
        note = document.add_paragraph()
        set_paragraph_text(note, "Controle rigoroso de administração de fármacos, horários e posologias recomendadas.", color="718096", size=9)

        add_medication_section("Histórico Detalhado")
        history_table = document.add_table(rows=1, cols=len(layout.history_columns))
        history_table.style = "Table Grid"
        for idx, label in enumerate(layout.history_columns):
            cell = history_table.rows[0].cells[idx]
            set_cell_shading(cell, "2B6CB0")
            set_cell_border(cell, "2B6CB0", "6")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="FFFFFF", size=7)

        for row_index, row in enumerate(layout.history_rows, start=1):
            cells = history_table.add_row().cells
            for idx, cell in enumerate(cells):
                set_cell_border(cell, "E2E8F0", "4")
                if row_index % 2 == 0:
                    set_cell_shading(cell, "F7FAFC")
            set_paragraph_text(cells[0].paragraphs[0], row["date"], bold=True, color="2D3748", size=8)
            set_paragraph_text(cells[1].paragraphs[0], row["time"], color="2D3748", size=8)
            set_paragraph_text(cells[2].paragraphs[0], row["medication_name"], bold=True, color="1A365D", size=9)
            dose_paragraph = cells[2].add_paragraph()
            set_paragraph_text(dose_paragraph, row["dose"] or "Dose não informada", color="4A5568", size=7)
            set_cell_shading(cells[3], "E2E8F0")
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_text(cells[3].paragraphs[0], row["quantity"], bold=True, color="2D3748", size=8)
            set_paragraph_text(cells[4].paragraphs[0], row["recurrence"], color="2D3748", size=8)
            set_paragraph_text(cells[5].paragraphs[0], row["observation"] or "", color="4A5568", size=7)

        add_medication_section("Informações Técnicas")
        technical_table = document.add_table(rows=1, cols=2)
        technical_table.style = "Table Grid"
        bar, content = technical_table.rows[0].cells
        set_cell_shading(bar, "CBD5E0")
        set_cell_border(bar, "CBD5E0", "4")
        set_cell_shading(content, "F7FAFC")
        set_cell_border(content, "F7FAFC", "4")
        for note_text in layout.technical_notes:
            paragraph = content.add_paragraph()
            set_paragraph_text(paragraph, note_text, color="4A5568", size=8)

        buffer = BytesIO()
        document.save(buffer)
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'docx')}\""
        return response

    if is_progress_export(rows, meta):
        layout = build_progress_export_layout(rows, meta)
        document = Document()
        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

        header_table = document.add_table(rows=1, cols=2)
        header_table.style = "Table Grid"
        for cell in header_table.rows[0].cells:
            set_cell_border(cell, "FFFFFF", "0")
        left, right = header_table.rows[0].cells
        set_paragraph_text(left.paragraphs[0], layout.title, bold=True, color="1A365D", size=21)
        subtitle = left.add_paragraph()
        set_paragraph_text(subtitle, "Cuidar Juntos", color="4A5568", size=10)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_text(right.paragraphs[0], f"Geração: {layout.generated_at}", color="718096", size=9)

        divider = document.add_table(rows=1, cols=1)
        divider.style = "Table Grid"
        set_cell_shading(divider.rows[0].cells[0], "E2E8F0")
        set_cell_border(divider.rows[0].cells[0], "E2E8F0", "4")

        document.add_paragraph()
        info_table = document.add_table(rows=1, cols=3)
        info_table.style = "Table Grid"
        for idx, (label, value) in enumerate(
            [
                ("PACIENTE", layout.patient_label),
                ("PROFISSIONAL RESPONSÁVEL", layout.professional_label),
                ("PERÍODO ANALISADO", layout.period_label),
            ]
        ):
            cell = info_table.rows[0].cells[idx]
            set_cell_border(cell, "EDF2F7", "8")
            set_cell_shading(cell, "F7FAFC")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="718096", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value or "Não informado", bold=True, color="2D3748", size=10)

        def add_progress_section(text: str) -> None:
            paragraph = document.add_paragraph()
            set_paragraph_text(paragraph, text, bold=True, color="2B6CB0", size=14)
            line = document.add_table(rows=1, cols=1)
            line.style = "Table Grid"
            set_cell_shading(line.rows[0].cells[0], "E2E8F0")
            set_cell_border(line.rows[0].cells[0], "E2E8F0", "4")

        add_progress_section("Resumo de Quadros Observados")
        summary_table = document.add_table(rows=1, cols=4)
        summary_table.style = "Table Grid"
        summary_colors = {"evolution": "22543D", "regression": "9B2C2C", "neutral": "4A5568", "main": "2B6CB0"}
        for idx, (label, value, kind) in enumerate(layout.summary_cards):
            cell = summary_table.rows[0].cells[idx]
            set_cell_border(cell, "BEE3F8" if kind == "main" else "E2E8F0", "8")
            set_cell_shading(cell, "EBF8FF" if kind == "main" else "FFFFFF")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="4A5568", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(
                paragraph,
                value,
                bold=True,
                color=summary_colors.get(kind, "2B6CB0"),
                size=17 if kind == "main" else 16,
            )
        note = document.add_paragraph()
        set_paragraph_text(note, "Consolidado do comportamento, autonomia e parâmetros de recuperação neurológica ou motora.", color="718096", size=9)

        add_progress_section("Histórico Detalhado")
        history_table = document.add_table(rows=1, cols=len(layout.history_columns))
        history_table.style = "Table Grid"
        for idx, label in enumerate(layout.history_columns):
            cell = history_table.rows[0].cells[idx]
            set_cell_shading(cell, "2B6CB0")
            set_cell_border(cell, "2B6CB0", "6")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="FFFFFF", size=7)

        tag_styles = {
            "evolution": ("C6F6D5", "22543D"),
            "regression": ("FED7D7", "9B2C2C"),
            "other": ("E2E8F0", "4A5568"),
        }
        for row_index, row in enumerate(layout.history_rows, start=1):
            cells = history_table.add_row().cells
            values = [row["date"], row["time"], row["classification"], row["recurrence"], row["observation"]]
            for idx, value in enumerate(values):
                set_cell_border(cells[idx], "E2E8F0", "4")
                if row_index % 2 == 0:
                    set_cell_shading(cells[idx], "F7FAFC")
                paragraph = cells[idx].paragraphs[0]
                if idx == 0:
                    set_paragraph_text(paragraph, value, bold=True, color="2D3748", size=8)
                elif idx == 2:
                    fill, color = tag_styles.get(row["classification_kind"], tag_styles["other"])
                    set_cell_shading(cells[idx], fill)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_text(paragraph, value.upper(), bold=True, color=color, size=7)
                elif idx == 4:
                    set_paragraph_text(paragraph, value or "", color="4A5568", size=7)
                else:
                    set_paragraph_text(paragraph, value or "Não informado", color="2D3748", size=8)

        add_progress_section("Informações Técnicas")
        technical_table = document.add_table(rows=1, cols=2)
        technical_table.style = "Table Grid"
        bar, content = technical_table.rows[0].cells
        set_cell_shading(bar, "CBD5E0")
        set_cell_border(bar, "CBD5E0", "4")
        set_cell_shading(content, "F7FAFC")
        set_cell_border(content, "F7FAFC", "4")
        for note_text in layout.technical_notes:
            paragraph = content.add_paragraph()
            set_paragraph_text(paragraph, note_text, color="4A5568", size=8)

        buffer = BytesIO()
        document.save(buffer)
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'docx')}\""
        return response

    if is_meal_export(rows, meta):
        layout = build_meal_export_layout(rows, meta)
        document = Document()
        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

        header_table = document.add_table(rows=1, cols=2)
        header_table.style = "Table Grid"
        for cell in header_table.rows[0].cells:
            set_cell_border(cell, "FFFFFF", "0")
        left, right = header_table.rows[0].cells
        set_paragraph_text(left.paragraphs[0], layout.title, bold=True, color="1A365D", size=21)
        subtitle = left.add_paragraph()
        set_paragraph_text(subtitle, "Cuidar Juntos", color="4A5568", size=10)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_text(right.paragraphs[0], f"Geração: {layout.generated_at}", color="718096", size=9)

        divider = document.add_table(rows=1, cols=1)
        divider.style = "Table Grid"
        set_cell_shading(divider.rows[0].cells[0], "E2E8F0")
        set_cell_border(divider.rows[0].cells[0], "E2E8F0", "4")

        document.add_paragraph()
        info_table = document.add_table(rows=1, cols=3)
        info_table.style = "Table Grid"
        for idx, (label, value) in enumerate(
            [
                ("PACIENTE", layout.patient_label),
                ("PROFISSIONAL RESPONSÁVEL", layout.professional_label),
                ("PERÍODO ANALISADO", layout.period_label),
            ]
        ):
            cell = info_table.rows[0].cells[idx]
            set_cell_border(cell, "EDF2F7", "8")
            set_cell_shading(cell, "F7FAFC")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="718096", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value or "Não informado", bold=True, color="2D3748", size=10)

        def add_meal_section(text: str) -> None:
            paragraph = document.add_paragraph()
            set_paragraph_text(paragraph, text, bold=True, color="2B6CB0", size=14)
            line = document.add_table(rows=1, cols=1)
            line.style = "Table Grid"
            set_cell_shading(line.rows[0].cells[0], "E2E8F0")
            set_cell_border(line.rows[0].cells[0], "E2E8F0", "4")

        add_meal_section("Resumo Nutricional")
        summary_table = document.add_table(rows=1, cols=4)
        summary_table.style = "Table Grid"
        summary_colors = {"good": "22543D", "poor": "9B2C2C", "blue": "2B6CB0", "main": "2B6CB0"}
        for idx, (label, value, kind) in enumerate(layout.summary_cards):
            cell = summary_table.rows[0].cells[idx]
            set_cell_border(cell, "BEE3F8" if kind == "main" else "E2E8F0", "8")
            set_cell_shading(cell, "EBF8FF" if kind == "main" else "FFFFFF")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="4A5568", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(
                paragraph,
                value,
                bold=True,
                color=summary_colors.get(kind, "2B6CB0"),
                size=17 if kind == "main" else 16,
            )
        note = document.add_paragraph()
        set_paragraph_text(note, "Acompanhamento diário do apetite, consistência alimentar e resposta a dietas específicas.", color="718096", size=9)

        add_meal_section("Histórico Detalhado")
        history_table = document.add_table(rows=1, cols=len(layout.history_columns))
        history_table.style = "Table Grid"
        for idx, label in enumerate(layout.history_columns):
            cell = history_table.rows[0].cells[idx]
            set_cell_shading(cell, "2B6CB0")
            set_cell_border(cell, "2B6CB0", "6")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="FFFFFF", size=7)

        tag_styles = {
            "good": ("C6F6D5", "22543D"),
            "poor": ("FED7D7", "9B2C2C"),
            "other": ("E2E8F0", "4A5568"),
        }
        for row_index, row in enumerate(layout.history_rows, start=1):
            cells = history_table.add_row().cells
            values = [row["date"], row["time"], row["meal_name"], row["recurrence"], row["acceptance"], row["observation"]]
            for idx, value in enumerate(values):
                set_cell_border(cells[idx], "E2E8F0", "4")
                if row_index % 2 == 0:
                    set_cell_shading(cells[idx], "F7FAFC")
                paragraph = cells[idx].paragraphs[0]
                if idx == 0:
                    set_paragraph_text(paragraph, value, bold=True, color="2D3748", size=8)
                elif idx == 2:
                    set_paragraph_text(paragraph, value, bold=True, color="1A365D", size=9)
                elif idx == 4:
                    fill, color = tag_styles.get(row["acceptance_kind"], tag_styles["other"])
                    set_cell_shading(cells[idx], fill)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_text(paragraph, value.upper(), bold=True, color=color, size=7)
                elif idx == 5:
                    set_paragraph_text(paragraph, value or "", color="4A5568", size=7)
                else:
                    set_paragraph_text(paragraph, value or "Não informado", color="2D3748", size=8)

        add_meal_section("Informações Técnicas")
        technical_table = document.add_table(rows=1, cols=2)
        technical_table.style = "Table Grid"
        bar, content = technical_table.rows[0].cells
        set_cell_shading(bar, "CBD5E0")
        set_cell_border(bar, "CBD5E0", "4")
        set_cell_shading(content, "F7FAFC")
        set_cell_border(content, "F7FAFC", "4")
        for note_text in layout.technical_notes:
            paragraph = content.add_paragraph()
            set_paragraph_text(paragraph, note_text, color="4A5568", size=8)

        buffer = BytesIO()
        document.save(buffer)
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'docx')}\""
        return response

    if is_activity_export(rows, meta):
        layout = build_activity_export_layout(rows, meta)
        document = Document()
        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

        header_table = document.add_table(rows=1, cols=2)
        header_table.style = "Table Grid"
        for cell in header_table.rows[0].cells:
            set_cell_border(cell, "FFFFFF", "0")
        left, right = header_table.rows[0].cells
        set_paragraph_text(left.paragraphs[0], layout.title, bold=True, color="1A365D", size=21)
        subtitle = left.add_paragraph()
        set_paragraph_text(subtitle, "Cuidar Juntos", color="4A5568", size=10)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_text(right.paragraphs[0], f"Geração: {layout.generated_at}", color="718096", size=9)

        divider = document.add_table(rows=1, cols=1)
        divider.style = "Table Grid"
        set_cell_shading(divider.rows[0].cells[0], "E2E8F0")
        set_cell_border(divider.rows[0].cells[0], "E2E8F0", "4")

        document.add_paragraph()
        info_table = document.add_table(rows=1, cols=3)
        info_table.style = "Table Grid"
        for idx, (label, value) in enumerate(
            [
                ("PACIENTE", layout.patient_label),
                ("PROFISSIONAL RESPONSÁVEL", layout.professional_label),
                ("PERÍODO ANALISADO", layout.period_label),
            ]
        ):
            cell = info_table.rows[0].cells[idx]
            set_cell_border(cell, "EDF2F7", "8")
            set_cell_shading(cell, "F7FAFC")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="718096", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, value or "Não informado", bold=True, color="2D3748", size=10)

        def add_activity_section(text: str) -> None:
            paragraph = document.add_paragraph()
            set_paragraph_text(paragraph, text, bold=True, color="2B6CB0", size=14)
            line = document.add_table(rows=1, cols=1)
            line.style = "Table Grid"
            set_cell_shading(line.rows[0].cells[0], "E2E8F0")
            set_cell_border(line.rows[0].cells[0], "E2E8F0", "4")

        add_activity_section("Resumo do Período")
        summary_table = document.add_table(rows=1, cols=4)
        summary_table.style = "Table Grid"
        for idx, (label, value, kind) in enumerate(layout.summary_cards):
            cell = summary_table.rows[0].cells[idx]
            set_cell_border(cell, "BEE3F8" if kind == "main" else "E2E8F0", "8")
            set_cell_shading(cell, "EBF8FF" if kind == "main" else "FFFFFF")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="4A5568", size=7)
            paragraph = cell.add_paragraph()
            set_paragraph_text(
                paragraph,
                value,
                bold=True,
                color="4A5568" if kind == "neutral" else "2B6CB0",
                size=17 if kind == "main" else 16,
            )
        note = document.add_paragraph()
        set_paragraph_text(note, "Registro e evolução do plano de mobilidade, condicionamento e reabilitação motora.", color="718096", size=9)

        add_activity_section("Histórico Detalhado")
        history_table = document.add_table(rows=1, cols=len(layout.history_columns))
        history_table.style = "Table Grid"
        for idx, label in enumerate(layout.history_columns):
            cell = history_table.rows[0].cells[idx]
            set_cell_shading(cell, "2B6CB0")
            set_cell_border(cell, "2B6CB0", "6")
            set_paragraph_text(cell.paragraphs[0], label, bold=True, color="FFFFFF", size=7)

        for row_index, row in enumerate(layout.history_rows, start=1):
            cells = history_table.add_row().cells
            values = [row["date"], row["time"], row["activity_name"], row["recurrence"], row["observation"]]
            for idx, value in enumerate(values):
                set_cell_border(cells[idx], "E2E8F0", "4")
                if row_index % 2 == 0:
                    set_cell_shading(cells[idx], "F7FAFC")
                paragraph = cells[idx].paragraphs[0]
                if idx == 0:
                    set_paragraph_text(paragraph, value, bold=True, color="2D3748", size=8)
                elif idx == 2:
                    set_paragraph_text(paragraph, value, bold=True, color="1A365D", size=9)
                elif idx == 4:
                    set_paragraph_text(paragraph, value or "", color="4A5568", size=7)
                else:
                    set_paragraph_text(paragraph, value or "Não informado", color="2D3748", size=8)

        add_activity_section("Informações Técnicas")
        technical_table = document.add_table(rows=1, cols=2)
        technical_table.style = "Table Grid"
        bar, content = technical_table.rows[0].cells
        set_cell_shading(bar, "CBD5E0")
        set_cell_border(bar, "CBD5E0", "4")
        set_cell_shading(content, "F7FAFC")
        set_cell_border(content, "F7FAFC", "4")
        for note_text in layout.technical_notes:
            paragraph = content.add_paragraph()
            set_paragraph_text(paragraph, note_text, color="4A5568", size=8)

        buffer = BytesIO()
        document.save(buffer)
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'docx')}\""
        return response

    document = Document()
    system = document.add_paragraph()
    system_run = system.add_run("Cuidar Juntos")
    system_run.bold = True
    title = document.add_heading(DOCUMENT_TITLE, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    summary_table = document.add_table(rows=0, cols=2)
    summary_table.autofit = False
    summary_table.style = "Table Grid"
    for label, value in meta.summary_rows():
        cells = summary_table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        if cells[0].paragraphs and cells[0].paragraphs[0].runs:
            cells[0].paragraphs[0].runs[0].bold = True

    document.add_paragraph()
    table = document.add_table(rows=1, cols=len(columns))
    table.autofit = False
    table.style = "Table Grid"
    col_widths_px = _column_width_px(columns)
    col_widths_in = [Inches(px / 96) for px in col_widths_px]
    for idx, width in enumerate(col_widths_in):
        table.columns[idx].width = width
    header_cells = table.rows[0].cells
    for idx, (_, label) in enumerate(columns):
        header_cells[idx].text = label
        if header_cells[idx].paragraphs and header_cells[idx].paragraphs[0].runs:
            header_cells[idx].paragraphs[0].runs[0].bold = True
        if header_cells[idx].paragraphs:
            header_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[idx].width = col_widths_in[idx]

    for row in rows:
        cells = table.add_row().cells
        for idx, (key, _) in enumerate(columns):
            cells[idx].text = row[key]
            cells[idx].width = col_widths_in[idx]
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    document.save(buffer)
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'docx')}\""
    return response


def export_as_pdf(
    rows: list[dict[str, str]],
    meta: ExportMetadata,
    columns: Sequence[tuple[str, str]] = COLUMNS,
) -> HttpResponse:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        return _export_pdf_inline(rows, meta, columns=columns)

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, title="Registros exportados")
    base_styles = getSampleStyleSheet()
    story: list[object] = []

    grid_gray = colors.HexColor("#D1D5DB")
    header_bg = colors.HexColor("#F3F4F6")

    header_style = ParagraphStyle(
        "TableHeader",
        parent=base_styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=8,
        leading=10,
        alignment=1,  # center
        wordWrap="CJK",
    )
    title_style = ParagraphStyle(
        "ExportTitle",
        parent=base_styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        spaceAfter=6,
        alignment=0,
    )
    subtitle_style = ParagraphStyle(
        "ExportSubtitle",
        parent=base_styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#374151"),
        spaceAfter=10,
    )
    cell_style = ParagraphStyle(
        "TableCell",
        parent=base_styles["Normal"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        alignment=0,  # left
        wordWrap="CJK",
    )

    if is_sleep_export(rows, meta):
        layout = build_sleep_export_layout(rows, meta)
        blue_dark = colors.HexColor("#123A5A")
        blue = colors.HexColor("#1D5F91")
        blue_light = colors.HexColor("#B7D7F2")
        gray_line = colors.HexColor("#D9E2EC")
        gray_text = colors.HexColor("#4B5563")
        header_gray = colors.HexColor("#F3F4F6")
        red = colors.HexColor("#B91C1C")
        green = colors.HexColor("#047857")
        amber = colors.HexColor("#B45309")

        section_style = ParagraphStyle(
            "SleepSection",
            parent=base_styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=blue,
            spaceBefore=18,
            spaceAfter=5,
        )
        sleep_title_style = ParagraphStyle(
            "SleepTitle",
            parent=base_styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=28,
            textColor=blue_dark,
            alignment=0,
            spaceAfter=1,
        )
        sleep_subtitle_style = ParagraphStyle(
            "SleepSubtitle",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            leading=12,
            textColor=gray_text,
        )
        generation_style = ParagraphStyle(
            "SleepGeneration",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            alignment=2,
            textColor=gray_text,
        )
        card_label_style = ParagraphStyle(
            "SleepCardLabel",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=colors.HexColor("#6B7280"),
            alignment=1,
        )
        card_value_style = ParagraphStyle(
            "SleepCardValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=blue,
            alignment=1,
        )
        info_value_style = ParagraphStyle(
            "SleepInfoValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#111827"),
            alignment=1,
            wordWrap="CJK",
        )
        small_style = ParagraphStyle(
            "SleepSmall",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=gray_text,
            wordWrap="CJK",
        )
        table_header_style = ParagraphStyle(
            "SleepTableHeader",
            parent=header_style,
            fontSize=7,
            leading=9,
            textColor=colors.HexColor("#6B7280"),
        )
        table_cell_style = ParagraphStyle(
            "SleepTableCell",
            parent=cell_style,
            fontSize=7,
            leading=9,
            textColor=colors.HexColor("#374151"),
            wordWrap="CJK",
        )
        bold_cell_style = ParagraphStyle("SleepBoldCell", parent=table_cell_style, fontName="Helvetica-Bold")
        missing_style = ParagraphStyle(
            "SleepMissing",
            parent=table_cell_style,
            fontName="Helvetica-Oblique",
            textColor=colors.HexColor("#6B7280"),
        )
        incomplete_style = ParagraphStyle(
            "SleepIncomplete",
            parent=table_cell_style,
            fontName="Helvetica-Oblique",
            textColor=red,
        )

        header_table = Table(
            [
                [
                    [Paragraph(layout.title, sleep_title_style), Paragraph("Cuidar Juntos", sleep_subtitle_style)],
                    Paragraph(f"Geração: {layout.generated_at}", generation_style),
                ]
            ],
            colWidths=[doc.width * 0.62, doc.width * 0.38],
        )
        header_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(header_table)
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=2, spaceAfter=18))

        info_data = [
            [
                [Paragraph("PACIENTE", card_label_style), Paragraph(layout.patient_label, info_value_style)],
                [Paragraph("PROFISSIONAL RESPONSÁVEL", card_label_style), Paragraph(layout.professional_label, info_value_style)],
                [Paragraph("PERÍODO ANALISADO", card_label_style), Paragraph(layout.period_label, info_value_style)],
            ]
        ]
        info_table = Table(info_data, colWidths=[doc.width / 3] * 3)
        info_table.setStyle(
            TableStyle(
                [
                    ("BOX", (0, 0), (-1, -1), 0.8, gray_line),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#EEF2F7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 12),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]
            )
        )
        story.append(info_table)

        story.append(Paragraph("Resumo do Sono", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=10))
        summary_data = [[
            [Paragraph(label, card_label_style), Paragraph(value, card_value_style)]
            for label, value in layout.summary_cards
        ]]
        summary_table = Table(summary_data, colWidths=[doc.width / 4] * 4)
        summary_table.setStyle(
            TableStyle(
                [
                    ("BOX", (0, 0), (0, 0), 1.0, blue_light),
                    ("BOX", (1, 0), (-1, -1), 0.7, gray_line),
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#F8FBFF")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(summary_table)
        story.append(
            Paragraph(
                (
                    f"Registros de início do sono: {layout.start_count} | "
                    f"Registros de término/despertar: {layout.end_count} | "
                    f"Total de registros avaliados: {layout.total_count}"
                ),
                small_style,
            )
        )

        story.append(Paragraph("Histórico Detalhado", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        history_data: list[list[object]] = [[Paragraph(label, table_header_style) for label in layout.history_columns]]
        if layout.history_rows:
            for row in layout.history_rows:
                status_color = {"done": green, "pending": amber, "missed": red}.get(row["status_key"], gray_text)
                status_style = ParagraphStyle(
                    f"SleepStatus{row['status_key'] or 'Other'}",
                    parent=table_cell_style,
                    fontName="Helvetica-Bold",
                    textColor=status_color,
                )
                status_parts = [Paragraph(row["status"] or "Não informado", status_style)]
                if row["caregiver"]:
                    status_parts.append(Paragraph(f"Resp: {row['caregiver']}", small_style))
                history_data.append(
                    [
                        Paragraph(row["date"], bold_cell_style),
                        Paragraph(row["start"] or "Não informado", table_cell_style if row["start"] else missing_style),
                        Paragraph(row["end"] or "Não informado", table_cell_style if row["end"] else missing_style),
                        Paragraph(row["duration"], bold_cell_style if row["duration_complete"] else incomplete_style),
                        status_parts,
                        Paragraph(row["observation"] or "", table_cell_style),
                    ]
                )
        else:
            history_data.append(
                [Paragraph("Nenhum registro de sono encontrado para os filtros selecionados.", table_cell_style)]
                + [Paragraph("Não informado", missing_style) for _ in layout.history_columns[1:]]
            )
        history_units = [11, 9, 9, 14, 19, 38]
        unit_total = sum(history_units) or 1
        history_widths = [(doc.width * unit / unit_total) for unit in history_units]
        history_table = Table(history_data, repeatRows=1, colWidths=history_widths)
        history_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), header_gray),
                    ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("LINEBELOW", (0, 0), (-1, -1), 0.5, gray_line),
                ]
            )
        )
        story.append(history_table)

        story.append(Paragraph("Informações Técnicas", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        technical_content = [[Paragraph(note, small_style)] for note in layout.technical_notes]
        technical_table = Table(
            [[Table([[""]], colWidths=[4]), Table(technical_content, colWidths=[doc.width - 28])]],
            colWidths=[8, doc.width - 8],
        )
        technical_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), blue_light),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(technical_table)

        doc.build(story)
        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'pdf')}\""
        return response

    if is_bathroom_export(rows, meta):
        layout = build_bathroom_export_layout(rows, meta)
        blue_dark = colors.HexColor("#123A5A")
        blue = colors.HexColor("#1D5F91")
        gray_line = colors.HexColor("#D9E2EC")
        gray_text = colors.HexColor("#4B5563")
        gray_soft = colors.HexColor("#F8FAFC")
        header_blue = colors.HexColor("#1D5F91")

        section_style = ParagraphStyle(
            "BathroomSection",
            parent=base_styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=blue,
            spaceBefore=18,
            spaceAfter=5,
        )
        bathroom_title_style = ParagraphStyle(
            "BathroomTitle",
            parent=base_styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=27,
            textColor=blue_dark,
            alignment=0,
            spaceAfter=1,
        )
        bathroom_subtitle_style = ParagraphStyle(
            "BathroomSubtitle",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            leading=12,
            textColor=gray_text,
        )
        generation_style = ParagraphStyle(
            "BathroomGeneration",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            alignment=2,
            textColor=gray_text,
        )
        label_style = ParagraphStyle(
            "BathroomLabel",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=colors.HexColor("#6B7280"),
            alignment=1,
        )
        info_value_style = ParagraphStyle(
            "BathroomInfoValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#111827"),
            alignment=1,
            wordWrap="CJK",
        )
        card_value_style = ParagraphStyle(
            "BathroomCardValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=blue,
            alignment=1,
        )
        small_style = ParagraphStyle(
            "BathroomSmall",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=gray_text,
            wordWrap="CJK",
        )
        table_header_style = ParagraphStyle(
            "BathroomTableHeader",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=colors.white,
            alignment=1,
        )
        table_cell_style = ParagraphStyle(
            "BathroomTableCell",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=7,
            leading=9,
            textColor=colors.HexColor("#374151"),
            wordWrap="CJK",
        )
        table_bold_style = ParagraphStyle("BathroomBoldCell", parent=table_cell_style, fontName="Helvetica-Bold")
        tag_styles = {
            "hygiene": (colors.HexColor("#DCEEFF"), colors.HexColor("#1D5F91")),
            "elimination": (colors.HexColor("#E5E7EB"), colors.HexColor("#374151")),
            "alert": (colors.HexColor("#FEE2E2"), colors.HexColor("#B91C1C")),
            "neutral": (colors.HexColor("#F3F4F6"), colors.HexColor("#4B5563")),
        }

        header_table = Table(
            [
                [
                    [Paragraph(layout.title, bathroom_title_style), Paragraph("Cuidar Juntos", bathroom_subtitle_style)],
                    Paragraph(f"Geração: {layout.generated_at}", generation_style),
                ]
            ],
            colWidths=[doc.width * 0.66, doc.width * 0.34],
        )
        header_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(header_table)
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=2, spaceAfter=18))

        info_table = Table(
            [
                [
                    [Paragraph("PACIENTE", label_style), Paragraph(layout.patient_label, info_value_style)],
                    [Paragraph("PROFISSIONAL RESPONSÁVEL", label_style), Paragraph(layout.professional_label, info_value_style)],
                    [Paragraph("PERÍODO ANALISADO", label_style), Paragraph(layout.period_label, info_value_style)],
                ]
            ],
            colWidths=[doc.width / 3] * 3,
        )
        info_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), gray_soft),
                    ("BOX", (0, 0), (-1, -1), 0.8, gray_line),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#EEF2F7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 12),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]
            )
        )
        story.append(info_table)

        story.append(Paragraph("Resumo das Ocorrências", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=10))
        summary_table = Table(
            [[
                [Paragraph(label, label_style), Paragraph(value, card_value_style)]
                for label, value in layout.summary_cards
            ]],
            colWidths=[doc.width / 4] * 4,
        )
        summary_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#EAF4FF")),
                    ("BOX", (0, 0), (-1, -1), 0.7, gray_line),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#EEF2F7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(summary_table)
        story.append(Paragraph("Monitoramento diário de rotinas de higiene fisiológica e cuidados básicos de saúde.", small_style))

        story.append(Paragraph("Histórico Detalhado", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        history_data: list[list[object]] = [[Paragraph(label, table_header_style) for label in layout.history_columns]]
        for row in layout.history_rows:
            tag_bg, tag_fg = tag_styles.get(row["tag_kind"], tag_styles["neutral"])
            tag_style = ParagraphStyle(
                f"BathroomTag{row['tag_kind']}",
                parent=table_header_style,
                textColor=tag_fg,
                fontSize=6,
                leading=8,
            )
            history_data.append(
                [
                    Paragraph(row["date"], table_bold_style),
                    Paragraph(row["time"], table_cell_style),
                    Paragraph(row["occurrence"].upper(), tag_style),
                    Paragraph(row["recurrence"], table_cell_style),
                    Paragraph(xml_escape(row["observation"] or "").replace("\n", "<br/>"), table_cell_style),
                ]
            )
        if not layout.history_rows:
            history_data.append(
                [Paragraph("Nenhum registro de banheiro encontrado para os filtros selecionados.", table_cell_style)]
                + [Paragraph("Não informado", table_cell_style) for _ in layout.history_columns[1:]]
            )
        history_units = [12, 10, 23, 17, 38]
        unit_total = sum(history_units)
        history_table = Table(
            history_data,
            repeatRows=1,
            colWidths=[doc.width * unit / unit_total for unit in history_units],
        )
        table_style_commands = [
            ("BACKGROUND", (0, 0), (-1, 0), header_blue),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, 0), (-1, -1), 0.5, gray_line),
        ]
        for idx, row in enumerate(layout.history_rows, start=1):
            if idx % 2 == 0:
                table_style_commands.append(("BACKGROUND", (0, idx), (-1, idx), gray_soft))
            tag_bg, _tag_fg = tag_styles.get(row["tag_kind"], tag_styles["neutral"])
            table_style_commands.append(("BACKGROUND", (2, idx), (2, idx), tag_bg))
        history_table.setStyle(TableStyle(table_style_commands))
        story.append(history_table)

        story.append(Paragraph("Informações Técnicas", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        technical_content = [[Paragraph(note, small_style)] for note in layout.technical_notes]
        technical_table = Table(
            [[Table([[""]], colWidths=[4]), Table(technical_content, colWidths=[doc.width - 28])]],
            colWidths=[8, doc.width - 8],
        )
        technical_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#D1D5DB")),
                    ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#F3F4F6")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(technical_table)

        doc.build(story)
        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'pdf')}\""
        return response

    if is_vital_export(rows, meta):
        layout = build_vital_export_layout(rows, meta)
        blue_dark = colors.HexColor("#1A365D")
        blue = colors.HexColor("#2B6CB0")
        gray_line = colors.HexColor("#E2E8F0")
        gray_text = colors.HexColor("#4A5568")
        gray_muted = colors.HexColor("#718096")
        gray_soft = colors.HexColor("#F7FAFC")

        section_style = ParagraphStyle(
            "VitalSection",
            parent=base_styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=blue,
            spaceBefore=18,
            spaceAfter=5,
        )
        vital_title_style = ParagraphStyle(
            "VitalTitle",
            parent=base_styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=27,
            textColor=blue_dark,
            alignment=0,
            spaceAfter=1,
        )
        vital_subtitle_style = ParagraphStyle(
            "VitalSubtitle",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            leading=12,
            textColor=gray_text,
        )
        generation_style = ParagraphStyle(
            "VitalGeneration",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            alignment=2,
            textColor=gray_muted,
        )
        label_style = ParagraphStyle(
            "VitalLabel",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=gray_muted,
            alignment=1,
        )
        info_value_style = ParagraphStyle(
            "VitalInfoValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#2D3748"),
            alignment=1,
            wordWrap="CJK",
        )
        card_value_base = ParagraphStyle(
            "VitalCardValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            alignment=1,
        )
        small_style = ParagraphStyle(
            "VitalSmall",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=gray_muted,
            wordWrap="CJK",
        )
        table_header_style = ParagraphStyle(
            "VitalTableHeader",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=colors.white,
            alignment=1,
        )
        table_cell_style = ParagraphStyle(
            "VitalTableCell",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=7,
            leading=9,
            textColor=colors.HexColor("#2D3748"),
            wordWrap="CJK",
        )
        table_bold_style = ParagraphStyle("VitalBoldCell", parent=table_cell_style, fontName="Helvetica-Bold")
        vital_name_style = ParagraphStyle("VitalNameCell", parent=table_cell_style, fontName="Helvetica-Bold", textColor=blue_dark)
        status_styles = {
            "normal": (colors.HexColor("#C6F6D5"), colors.HexColor("#22543D")),
            "warning": (colors.HexColor("#FEFCBF"), colors.HexColor("#744210")),
            "alert": (colors.HexColor("#FED7D7"), colors.HexColor("#9B2C2C")),
            "neutral": (colors.HexColor("#E2E8F0"), gray_text),
        }
        card_value_colors = {
            "main": blue,
            "normal": colors.HexColor("#22543D"),
            "warning": colors.HexColor("#B7791F"),
            "alert": colors.HexColor("#9B2C2C"),
        }

        header_table = Table(
            [
                [
                    [Paragraph(layout.title, vital_title_style), Paragraph("Cuidar Juntos", vital_subtitle_style)],
                    Paragraph(f"Geração: {layout.generated_at}", generation_style),
                ]
            ],
            colWidths=[doc.width * 0.66, doc.width * 0.34],
        )
        header_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(header_table)
        story.append(HRFlowable(width="100%", thickness=1.0, color=gray_line, spaceBefore=2, spaceAfter=18))

        info_table = Table(
            [
                [
                    [Paragraph("PACIENTE", label_style), Paragraph(layout.patient_label, info_value_style)],
                    [Paragraph("PROFISSIONAL RESPONSÁVEL", label_style), Paragraph(layout.professional_label, info_value_style)],
                    [Paragraph("PERÍODO ANALISADO", label_style), Paragraph(layout.period_label, info_value_style)],
                ]
            ],
            colWidths=[doc.width / 3] * 3,
        )
        info_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), gray_soft),
                    ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#EDF2F7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 12),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]
            )
        )
        story.append(info_table)

        story.append(Paragraph("Resumo de Alertas Clínicos", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=10))
        summary_cells = []
        for label, value, kind in layout.summary_cards:
            value_style = ParagraphStyle(
                f"VitalCard{kind}",
                parent=card_value_base,
                fontSize=18 if kind == "main" else 16,
                textColor=card_value_colors[kind],
            )
            summary_cells.append([Paragraph(label, label_style), Paragraph(value, value_style)])
        summary_table = Table([summary_cells], colWidths=[doc.width / 4] * 4)
        summary_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#EBF8FF")),
                    ("BOX", (0, 0), (0, 0), 0.8, colors.HexColor("#BEE3F8")),
                    ("BOX", (1, 0), (-1, -1), 0.7, gray_line),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, gray_line),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(summary_table)
        story.append(Paragraph("Triagem e monitoramento preventivo de parâmetros hemodinâmicos e térmicos.", small_style))

        story.append(Paragraph("Histórico Detalhado", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        history_data: list[list[object]] = [[Paragraph(label, table_header_style) for label in layout.history_columns]]
        for row in layout.history_rows:
            tag_bg, tag_fg = status_styles.get(row["status_level"], status_styles["neutral"])
            tag_style = ParagraphStyle(
                f"VitalTag{row['status_level']}",
                parent=table_header_style,
                textColor=tag_fg,
                fontSize=6,
                leading=8,
            )
            history_data.append(
                [
                    Paragraph(row["date"], table_bold_style),
                    Paragraph(row["time"], table_cell_style),
                    Paragraph(row["vital"], vital_name_style),
                    Paragraph(row["recurrence"], table_cell_style),
                    Paragraph(row["status"].upper(), tag_style),
                    Paragraph(xml_escape(row["observation"] or "").replace("\n", "<br/>"), small_style),
                ]
            )
        if not layout.history_rows:
            history_data.append(
                [Paragraph("Nenhum registro de sinais vitais encontrado para os filtros selecionados.", table_cell_style)]
                + [Paragraph("Não informado", table_cell_style) for _ in layout.history_columns[1:]]
            )
        history_units = [11, 9, 23, 15, 18, 24]
        unit_total = sum(history_units)
        history_table = Table(
            history_data,
            repeatRows=1,
            colWidths=[doc.width * unit / unit_total for unit in history_units],
        )
        table_style_commands = [
            ("BACKGROUND", (0, 0), (-1, 0), blue),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, 0), (-1, -1), 0.5, gray_line),
        ]
        for idx, row in enumerate(layout.history_rows, start=1):
            if idx % 2 == 0:
                table_style_commands.append(("BACKGROUND", (0, idx), (-1, idx), gray_soft))
            tag_bg, _tag_fg = status_styles.get(row["status_level"], status_styles["neutral"])
            table_style_commands.append(("BACKGROUND", (4, idx), (4, idx), tag_bg))
        history_table.setStyle(TableStyle(table_style_commands))
        story.append(history_table)

        story.append(Paragraph("Informações Técnicas", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        technical_content = [[Paragraph(note, small_style)] for note in layout.technical_notes]
        technical_table = Table(
            [[Table([[""]], colWidths=[4]), Table(technical_content, colWidths=[doc.width - 28])]],
            colWidths=[8, doc.width - 8],
        )
        technical_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#CBD5E0")),
                    ("BACKGROUND", (1, 0), (1, 0), gray_soft),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(technical_table)

        doc.build(story)
        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'pdf')}\""
        return response

    if is_medication_export(rows, meta):
        layout = build_medication_export_layout(rows, meta)
        blue_dark = colors.HexColor("#1A365D")
        blue = colors.HexColor("#2B6CB0")
        gray_line = colors.HexColor("#E2E8F0")
        gray_text = colors.HexColor("#4A5568")
        gray_muted = colors.HexColor("#718096")
        gray_soft = colors.HexColor("#F7FAFC")

        section_style = ParagraphStyle(
            "MedicationSection",
            parent=base_styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=blue,
            spaceBefore=18,
            spaceAfter=5,
        )
        medication_title_style = ParagraphStyle(
            "MedicationTitle",
            parent=base_styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=27,
            textColor=blue_dark,
            alignment=0,
            spaceAfter=1,
        )
        medication_subtitle_style = ParagraphStyle(
            "MedicationSubtitle",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            leading=12,
            textColor=gray_text,
        )
        generation_style = ParagraphStyle(
            "MedicationGeneration",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            alignment=2,
            textColor=gray_muted,
        )
        label_style = ParagraphStyle(
            "MedicationLabel",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=gray_muted,
            alignment=1,
        )
        info_value_style = ParagraphStyle(
            "MedicationInfoValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#2D3748"),
            alignment=1,
            wordWrap="CJK",
        )
        card_value_base = ParagraphStyle(
            "MedicationCardValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            alignment=1,
            textColor=blue,
        )
        small_style = ParagraphStyle(
            "MedicationSmall",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=gray_muted,
            wordWrap="CJK",
        )
        table_header_style = ParagraphStyle(
            "MedicationTableHeader",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=colors.white,
            alignment=1,
        )
        table_cell_style = ParagraphStyle(
            "MedicationTableCell",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=7,
            leading=9,
            textColor=colors.HexColor("#2D3748"),
            wordWrap="CJK",
        )
        table_bold_style = ParagraphStyle("MedicationBoldCell", parent=table_cell_style, fontName="Helvetica-Bold")
        med_name_style = ParagraphStyle("MedicationNameCell", parent=table_cell_style, fontName="Helvetica-Bold", fontSize=8, textColor=blue_dark)
        med_dose_style = ParagraphStyle("MedicationDoseCell", parent=table_cell_style, fontSize=7, textColor=gray_text)
        badge_style = ParagraphStyle("MedicationBadge", parent=table_cell_style, fontName="Helvetica-Bold", alignment=1)

        header_table = Table(
            [
                [
                    [Paragraph(layout.title, medication_title_style), Paragraph("Cuidar Juntos", medication_subtitle_style)],
                    Paragraph(f"Geração: {layout.generated_at}", generation_style),
                ]
            ],
            colWidths=[doc.width * 0.66, doc.width * 0.34],
        )
        header_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(header_table)
        story.append(HRFlowable(width="100%", thickness=1.0, color=gray_line, spaceBefore=2, spaceAfter=18))

        info_table = Table(
            [
                [
                    [Paragraph("PACIENTE", label_style), Paragraph(layout.patient_label, info_value_style)],
                    [Paragraph("PROFISSIONAL RESPONSÁVEL", label_style), Paragraph(layout.professional_label, info_value_style)],
                    [Paragraph("PERÍODO ANALISADO", label_style), Paragraph(layout.period_label, info_value_style)],
                ]
            ],
            colWidths=[doc.width / 3] * 3,
        )
        info_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), gray_soft),
                    ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#EDF2F7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 12),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]
            )
        )
        story.append(info_table)

        story.append(Paragraph("Resumo de Ministrações", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=10))
        summary_cells = []
        for label, value, kind in layout.summary_cards:
            value_style = ParagraphStyle(
                f"MedicationCard{kind}",
                parent=card_value_base,
                fontSize=18 if kind == "main" else 16,
                textColor=gray_text if kind == "neutral" else blue,
            )
            summary_cells.append([Paragraph(label, label_style), Paragraph(value, value_style)])
        summary_table = Table([summary_cells], colWidths=[doc.width / 4] * 4)
        summary_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#EBF8FF")),
                    ("BOX", (0, 0), (0, 0), 0.8, colors.HexColor("#BEE3F8")),
                    ("BOX", (1, 0), (-1, -1), 0.7, gray_line),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, gray_line),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(summary_table)
        story.append(Paragraph("Controle rigoroso de administração de fármacos, horários e posologias recomendadas.", small_style))

        story.append(Paragraph("Histórico Detalhado", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        history_data: list[list[object]] = [[Paragraph(label, table_header_style) for label in layout.history_columns]]
        for row in layout.history_rows:
            med_cell = [Paragraph(row["medication_name"], med_name_style)]
            if row["dose"]:
                med_cell.append(Paragraph(row["dose"], med_dose_style))
            history_data.append(
                [
                    Paragraph(row["date"], table_bold_style),
                    Paragraph(row["time"], table_cell_style),
                    med_cell,
                    Paragraph(row["quantity"], badge_style),
                    Paragraph(row["recurrence"], table_cell_style),
                    Paragraph(xml_escape(row["observation"] or "").replace("\n", "<br/>"), small_style),
                ]
            )
        if not layout.history_rows:
            history_data.append(
                [Paragraph("Nenhum registro de medicação encontrado para os filtros selecionados.", table_cell_style)]
                + [Paragraph("Não informado", table_cell_style) for _ in layout.history_columns[1:]]
            )
        history_units = [11, 9, 28, 15, 15, 22]
        unit_total = sum(history_units)
        history_table = Table(
            history_data,
            repeatRows=1,
            colWidths=[doc.width * unit / unit_total for unit in history_units],
        )
        table_style_commands = [
            ("BACKGROUND", (0, 0), (-1, 0), blue),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, 0), (-1, -1), 0.5, gray_line),
        ]
        for idx, _row in enumerate(layout.history_rows, start=1):
            if idx % 2 == 0:
                table_style_commands.append(("BACKGROUND", (0, idx), (-1, idx), gray_soft))
            table_style_commands.append(("BACKGROUND", (3, idx), (3, idx), colors.HexColor("#E2E8F0")))
        history_table.setStyle(TableStyle(table_style_commands))
        story.append(history_table)

        story.append(Paragraph("Informações Técnicas", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        technical_content = [[Paragraph(note, small_style)] for note in layout.technical_notes]
        technical_table = Table(
            [[Table([[""]], colWidths=[4]), Table(technical_content, colWidths=[doc.width - 28])]],
            colWidths=[8, doc.width - 8],
        )
        technical_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#CBD5E0")),
                    ("BACKGROUND", (1, 0), (1, 0), gray_soft),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(technical_table)

        doc.build(story)
        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'pdf')}\""
        return response

    if is_progress_export(rows, meta):
        layout = build_progress_export_layout(rows, meta)
        blue_dark = colors.HexColor("#1A365D")
        blue = colors.HexColor("#2B6CB0")
        gray_line = colors.HexColor("#E2E8F0")
        gray_text = colors.HexColor("#4A5568")
        gray_muted = colors.HexColor("#718096")
        gray_soft = colors.HexColor("#F7FAFC")
        green = colors.HexColor("#22543D")
        red = colors.HexColor("#9B2C2C")

        section_style = ParagraphStyle(
            "ProgressSection",
            parent=base_styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=blue,
            spaceBefore=18,
            spaceAfter=5,
        )
        progress_title_style = ParagraphStyle(
            "ProgressTitle",
            parent=base_styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=27,
            textColor=blue_dark,
            alignment=0,
            spaceAfter=1,
        )
        progress_subtitle_style = ParagraphStyle(
            "ProgressSubtitle",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=12,
            textColor=gray_text,
        )
        generation_style = ParagraphStyle(
            "ProgressGeneration",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            alignment=2,
            textColor=gray_muted,
        )
        label_style = ParagraphStyle(
            "ProgressLabel",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=gray_muted,
            alignment=1,
        )
        info_value_style = ParagraphStyle(
            "ProgressInfoValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#2D3748"),
            alignment=1,
            wordWrap="CJK",
        )
        card_value_base = ParagraphStyle(
            "ProgressCardValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            alignment=1,
            textColor=blue,
        )
        small_style = ParagraphStyle(
            "ProgressSmall",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=gray_text,
            wordWrap="CJK",
        )
        muted_small_style = ParagraphStyle("ProgressMutedSmall", parent=small_style, textColor=gray_muted)
        table_header_style = ParagraphStyle(
            "ProgressTableHeader",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=colors.white,
            alignment=1,
        )
        table_cell_style = ParagraphStyle(
            "ProgressTableCell",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=7,
            leading=9,
            textColor=colors.HexColor("#2D3748"),
            wordWrap="CJK",
        )
        table_bold_style = ParagraphStyle("ProgressBoldCell", parent=table_cell_style, fontName="Helvetica-Bold")
        tag_style_base = ParagraphStyle("ProgressTag", parent=table_header_style, fontSize=6, leading=8)

        header_table = Table(
            [
                [
                    [Paragraph(layout.title, progress_title_style), Paragraph("Cuidar Juntos", progress_subtitle_style)],
                    Paragraph(f"Geração: {layout.generated_at}", generation_style),
                ]
            ],
            colWidths=[doc.width * 0.66, doc.width * 0.34],
        )
        header_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(header_table)
        story.append(HRFlowable(width="100%", thickness=1.0, color=gray_line, spaceBefore=2, spaceAfter=18))

        info_table = Table(
            [
                [
                    [Paragraph("PACIENTE", label_style), Paragraph(layout.patient_label, info_value_style)],
                    [Paragraph("PROFISSIONAL RESPONSÁVEL", label_style), Paragraph(layout.professional_label, info_value_style)],
                    [Paragraph("PERÍODO ANALISADO", label_style), Paragraph(layout.period_label, info_value_style)],
                ]
            ],
            colWidths=[doc.width / 3] * 3,
        )
        info_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), gray_soft),
                    ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#EDF2F7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 12),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]
            )
        )
        story.append(info_table)

        story.append(Paragraph("Resumo de Quadros Observados", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=10))
        summary_cells = []
        summary_colors = {"evolution": green, "regression": red, "neutral": gray_text, "main": blue}
        for label, value, kind in layout.summary_cards:
            value_style = ParagraphStyle(
                f"ProgressCard{kind}",
                parent=card_value_base,
                fontSize=18 if kind == "main" else 16,
                textColor=summary_colors.get(kind, blue),
            )
            summary_cells.append([Paragraph(label, label_style), Paragraph(value, value_style)])
        summary_table = Table([summary_cells], colWidths=[doc.width / 4] * 4)
        summary_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#EBF8FF")),
                    ("BOX", (0, 0), (0, 0), 0.8, colors.HexColor("#BEE3F8")),
                    ("BOX", (1, 0), (-1, -1), 0.7, gray_line),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, gray_line),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(summary_table)
        story.append(Paragraph("Consolidado do comportamento, autonomia e parâmetros de recuperação neurológica ou motora.", muted_small_style))

        story.append(Paragraph("Histórico Detalhado", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        history_data: list[list[object]] = [[Paragraph(label, table_header_style) for label in layout.history_columns]]
        tag_styles = {
            "evolution": (colors.HexColor("#C6F6D5"), green),
            "regression": (colors.HexColor("#FED7D7"), red),
            "other": (colors.HexColor("#E2E8F0"), gray_text),
        }
        for row in layout.history_rows:
            _tag_bg, tag_fg = tag_styles.get(row["classification_kind"], tag_styles["other"])
            tag_style = ParagraphStyle(
                f"ProgressTag{row['classification_kind']}",
                parent=tag_style_base,
                textColor=tag_fg,
            )
            history_data.append(
                [
                    Paragraph(row["date"], table_bold_style),
                    Paragraph(row["time"], table_cell_style),
                    Paragraph(row["classification"].upper(), tag_style),
                    Paragraph(row["recurrence"], table_cell_style),
                    Paragraph(xml_escape(row["observation"] or "").replace("\n", "<br/>"), small_style),
                ]
            )
        if not layout.history_rows:
            history_data.append(
                [Paragraph("Nenhum registro de evolução/regressão encontrado para os filtros selecionados.", table_cell_style)]
                + [Paragraph("Não informado", table_cell_style) for _ in layout.history_columns[1:]]
            )
        history_units = [11, 9, 22, 16, 42]
        unit_total = sum(history_units)
        history_table = Table(
            history_data,
            repeatRows=1,
            colWidths=[doc.width * unit / unit_total for unit in history_units],
        )
        table_style_commands = [
            ("BACKGROUND", (0, 0), (-1, 0), blue),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("ALIGN", (2, 1), (2, -1), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, 0), (-1, -1), 0.5, gray_line),
        ]
        for idx, row in enumerate(layout.history_rows, start=1):
            if idx % 2 == 0:
                table_style_commands.append(("BACKGROUND", (0, idx), (-1, idx), gray_soft))
            tag_bg, _tag_fg = tag_styles.get(row["classification_kind"], tag_styles["other"])
            table_style_commands.append(("BACKGROUND", (2, idx), (2, idx), tag_bg))
        history_table.setStyle(TableStyle(table_style_commands))
        story.append(history_table)

        story.append(Paragraph("Informações Técnicas", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        technical_content = [[Paragraph(note, small_style)] for note in layout.technical_notes]
        technical_table = Table(
            [[Table([[""]], colWidths=[4]), Table(technical_content, colWidths=[doc.width - 28])]],
            colWidths=[8, doc.width - 8],
        )
        technical_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#CBD5E0")),
                    ("BACKGROUND", (1, 0), (1, 0), gray_soft),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(technical_table)

        doc.build(story)
        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'pdf')}\""
        return response

    if is_meal_export(rows, meta):
        layout = build_meal_export_layout(rows, meta)
        blue_dark = colors.HexColor("#1A365D")
        blue = colors.HexColor("#2B6CB0")
        gray_line = colors.HexColor("#E2E8F0")
        gray_text = colors.HexColor("#4A5568")
        gray_muted = colors.HexColor("#718096")
        gray_soft = colors.HexColor("#F7FAFC")
        green = colors.HexColor("#22543D")
        red = colors.HexColor("#9B2C2C")

        section_style = ParagraphStyle(
            "MealSection",
            parent=base_styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=blue,
            spaceBefore=18,
            spaceAfter=5,
        )
        meal_title_style = ParagraphStyle(
            "MealTitle",
            parent=base_styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=27,
            textColor=blue_dark,
            alignment=0,
            spaceAfter=1,
        )
        meal_subtitle_style = ParagraphStyle(
            "MealSubtitle",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=12,
            textColor=gray_text,
        )
        generation_style = ParagraphStyle(
            "MealGeneration",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            alignment=2,
            textColor=gray_muted,
        )
        label_style = ParagraphStyle(
            "MealLabel",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=gray_muted,
            alignment=1,
        )
        info_value_style = ParagraphStyle(
            "MealInfoValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#2D3748"),
            alignment=1,
            wordWrap="CJK",
        )
        card_value_base = ParagraphStyle(
            "MealCardValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            alignment=1,
            textColor=blue,
        )
        small_style = ParagraphStyle(
            "MealSmall",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=gray_text,
            wordWrap="CJK",
        )
        muted_small_style = ParagraphStyle("MealMutedSmall", parent=small_style, textColor=gray_muted)
        table_header_style = ParagraphStyle(
            "MealTableHeader",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=colors.white,
            alignment=1,
        )
        table_cell_style = ParagraphStyle(
            "MealTableCell",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=7,
            leading=9,
            textColor=colors.HexColor("#2D3748"),
            wordWrap="CJK",
        )
        table_bold_style = ParagraphStyle("MealBoldCell", parent=table_cell_style, fontName="Helvetica-Bold")
        meal_name_style = ParagraphStyle(
            "MealNameCell",
            parent=table_cell_style,
            fontName="Helvetica-Bold",
            fontSize=8,
            textColor=blue_dark,
        )
        tag_style_base = ParagraphStyle(
            "MealTag",
            parent=table_header_style,
            fontSize=6,
            leading=8,
        )

        header_table = Table(
            [
                [
                    [Paragraph(layout.title, meal_title_style), Paragraph("Cuidar Juntos", meal_subtitle_style)],
                    Paragraph(f"Geração: {layout.generated_at}", generation_style),
                ]
            ],
            colWidths=[doc.width * 0.66, doc.width * 0.34],
        )
        header_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(header_table)
        story.append(HRFlowable(width="100%", thickness=1.0, color=gray_line, spaceBefore=2, spaceAfter=18))

        info_table = Table(
            [
                [
                    [Paragraph("PACIENTE", label_style), Paragraph(layout.patient_label, info_value_style)],
                    [Paragraph("PROFISSIONAL RESPONSÁVEL", label_style), Paragraph(layout.professional_label, info_value_style)],
                    [Paragraph("PERÍODO ANALISADO", label_style), Paragraph(layout.period_label, info_value_style)],
                ]
            ],
            colWidths=[doc.width / 3] * 3,
        )
        info_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), gray_soft),
                    ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#EDF2F7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 12),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]
            )
        )
        story.append(info_table)

        story.append(Paragraph("Resumo Nutricional", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=10))
        summary_cells = []
        summary_colors = {"good": green, "poor": red, "blue": blue, "main": blue}
        for label, value, kind in layout.summary_cards:
            value_style = ParagraphStyle(
                f"MealCard{kind}",
                parent=card_value_base,
                fontSize=18 if kind == "main" else 16,
                textColor=summary_colors.get(kind, blue),
            )
            summary_cells.append([Paragraph(label, label_style), Paragraph(value, value_style)])
        summary_table = Table([summary_cells], colWidths=[doc.width / 4] * 4)
        summary_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#EBF8FF")),
                    ("BOX", (0, 0), (0, 0), 0.8, colors.HexColor("#BEE3F8")),
                    ("BOX", (1, 0), (-1, -1), 0.7, gray_line),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, gray_line),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(summary_table)
        story.append(Paragraph("Acompanhamento diário do apetite, consistência alimentar e resposta a dietas específicas.", muted_small_style))

        story.append(Paragraph("Histórico Detalhado", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        history_data: list[list[object]] = [[Paragraph(label, table_header_style) for label in layout.history_columns]]
        tag_styles = {
            "good": (colors.HexColor("#C6F6D5"), green),
            "poor": (colors.HexColor("#FED7D7"), red),
            "other": (colors.HexColor("#E2E8F0"), gray_text),
        }
        for row in layout.history_rows:
            _tag_bg, tag_fg = tag_styles.get(row["acceptance_kind"], tag_styles["other"])
            tag_style = ParagraphStyle(
                f"MealTag{row['acceptance_kind']}",
                parent=tag_style_base,
                textColor=tag_fg,
            )
            history_data.append(
                [
                    Paragraph(row["date"], table_bold_style),
                    Paragraph(row["time"], table_cell_style),
                    Paragraph(row["meal_name"], meal_name_style),
                    Paragraph(row["recurrence"], table_cell_style),
                    Paragraph(row["acceptance"].upper(), tag_style),
                    Paragraph(xml_escape(row["observation"] or "").replace("\n", "<br/>"), small_style),
                ]
            )
        if not layout.history_rows:
            history_data.append(
                [Paragraph("Nenhum registro de alimentação encontrado para os filtros selecionados.", table_cell_style)]
                + [Paragraph("Não informado", table_cell_style) for _ in layout.history_columns[1:]]
            )
        history_units = [11, 9, 25, 15, 18, 22]
        unit_total = sum(history_units)
        history_table = Table(
            history_data,
            repeatRows=1,
            colWidths=[doc.width * unit / unit_total for unit in history_units],
        )
        table_style_commands = [
            ("BACKGROUND", (0, 0), (-1, 0), blue),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("ALIGN", (4, 1), (4, -1), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, 0), (-1, -1), 0.5, gray_line),
        ]
        for idx, row in enumerate(layout.history_rows, start=1):
            if idx % 2 == 0:
                table_style_commands.append(("BACKGROUND", (0, idx), (-1, idx), gray_soft))
            tag_bg, _tag_fg = tag_styles.get(row["acceptance_kind"], tag_styles["other"])
            table_style_commands.append(("BACKGROUND", (4, idx), (4, idx), tag_bg))
        history_table.setStyle(TableStyle(table_style_commands))
        story.append(history_table)

        story.append(Paragraph("Informações Técnicas", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        technical_content = [[Paragraph(note, small_style)] for note in layout.technical_notes]
        technical_table = Table(
            [[Table([[""]], colWidths=[4]), Table(technical_content, colWidths=[doc.width - 28])]],
            colWidths=[8, doc.width - 8],
        )
        technical_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#CBD5E0")),
                    ("BACKGROUND", (1, 0), (1, 0), gray_soft),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(technical_table)

        doc.build(story)
        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'pdf')}\""
        return response

    if is_activity_export(rows, meta):
        layout = build_activity_export_layout(rows, meta)
        blue_dark = colors.HexColor("#1A365D")
        blue = colors.HexColor("#2B6CB0")
        gray_line = colors.HexColor("#E2E8F0")
        gray_text = colors.HexColor("#4A5568")
        gray_muted = colors.HexColor("#718096")
        gray_soft = colors.HexColor("#F7FAFC")

        section_style = ParagraphStyle(
            "ActivitySection",
            parent=base_styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=blue,
            spaceBefore=18,
            spaceAfter=5,
        )
        activity_title_style = ParagraphStyle(
            "ActivityTitle",
            parent=base_styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=27,
            textColor=blue_dark,
            alignment=0,
            spaceAfter=1,
        )
        activity_subtitle_style = ParagraphStyle(
            "ActivitySubtitle",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=12,
            textColor=gray_text,
        )
        generation_style = ParagraphStyle(
            "ActivityGeneration",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            alignment=2,
            textColor=gray_muted,
        )
        label_style = ParagraphStyle(
            "ActivityLabel",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=gray_muted,
            alignment=1,
        )
        info_value_style = ParagraphStyle(
            "ActivityInfoValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#2D3748"),
            alignment=1,
            wordWrap="CJK",
        )
        card_value_base = ParagraphStyle(
            "ActivityCardValue",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            alignment=1,
            textColor=blue,
        )
        small_style = ParagraphStyle(
            "ActivitySmall",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=gray_muted,
            wordWrap="CJK",
        )
        table_header_style = ParagraphStyle(
            "ActivityTableHeader",
            parent=base_styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=7,
            leading=9,
            textColor=colors.white,
            alignment=1,
        )
        table_cell_style = ParagraphStyle(
            "ActivityTableCell",
            parent=base_styles["Normal"],
            fontName="Helvetica",
            fontSize=7,
            leading=9,
            textColor=colors.HexColor("#2D3748"),
            wordWrap="CJK",
        )
        table_bold_style = ParagraphStyle("ActivityBoldCell", parent=table_cell_style, fontName="Helvetica-Bold")
        activity_name_style = ParagraphStyle(
            "ActivityNameCell",
            parent=table_cell_style,
            fontName="Helvetica-Bold",
            fontSize=8,
            textColor=blue_dark,
        )

        header_table = Table(
            [
                [
                    [Paragraph(layout.title, activity_title_style), Paragraph("Cuidar Juntos", activity_subtitle_style)],
                    Paragraph(f"Geração: {layout.generated_at}", generation_style),
                ]
            ],
            colWidths=[doc.width * 0.66, doc.width * 0.34],
        )
        header_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(header_table)
        story.append(HRFlowable(width="100%", thickness=1.0, color=gray_line, spaceBefore=2, spaceAfter=18))

        info_table = Table(
            [
                [
                    [Paragraph("PACIENTE", label_style), Paragraph(layout.patient_label, info_value_style)],
                    [Paragraph("PROFISSIONAL RESPONSÁVEL", label_style), Paragraph(layout.professional_label, info_value_style)],
                    [Paragraph("PERÍODO ANALISADO", label_style), Paragraph(layout.period_label, info_value_style)],
                ]
            ],
            colWidths=[doc.width / 3] * 3,
        )
        info_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), gray_soft),
                    ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#EDF2F7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 12),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]
            )
        )
        story.append(info_table)

        story.append(Paragraph("Resumo do Período", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=10))
        summary_cells = []
        for label, value, kind in layout.summary_cards:
            value_style = ParagraphStyle(
                f"ActivityCard{kind}",
                parent=card_value_base,
                fontSize=18 if kind == "main" else 16,
                textColor=gray_text if kind == "neutral" else blue,
            )
            summary_cells.append([Paragraph(label, label_style), Paragraph(value, value_style)])
        summary_table = Table([summary_cells], colWidths=[doc.width / 4] * 4)
        summary_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#EBF8FF")),
                    ("BOX", (0, 0), (0, 0), 0.8, colors.HexColor("#BEE3F8")),
                    ("BOX", (1, 0), (-1, -1), 0.7, gray_line),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, gray_line),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(summary_table)
        story.append(Paragraph("Registro e evolução do plano de mobilidade, condicionamento e reabilitação motora.", small_style))

        story.append(Paragraph("Histórico Detalhado", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        history_data: list[list[object]] = [[Paragraph(label, table_header_style) for label in layout.history_columns]]
        for row in layout.history_rows:
            history_data.append(
                [
                    Paragraph(row["date"], table_bold_style),
                    Paragraph(row["time"], table_cell_style),
                    Paragraph(row["activity_name"], activity_name_style),
                    Paragraph(row["recurrence"], table_cell_style),
                    Paragraph(xml_escape(row["observation"] or "").replace("\n", "<br/>"), small_style),
                ]
            )
        if not layout.history_rows:
            history_data.append(
                [Paragraph("Nenhum registro de atividade encontrado para os filtros selecionados.", table_cell_style)]
                + [Paragraph("Não informado", table_cell_style) for _ in layout.history_columns[1:]]
            )
        history_units = [11, 9, 30, 15, 35]
        unit_total = sum(history_units)
        history_table = Table(
            history_data,
            repeatRows=1,
            colWidths=[doc.width * unit / unit_total for unit in history_units],
        )
        table_style_commands = [
            ("BACKGROUND", (0, 0), (-1, 0), blue),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW", (0, 0), (-1, -1), 0.5, gray_line),
        ]
        for idx, _row in enumerate(layout.history_rows, start=1):
            if idx % 2 == 0:
                table_style_commands.append(("BACKGROUND", (0, idx), (-1, idx), gray_soft))
        history_table.setStyle(TableStyle(table_style_commands))
        story.append(history_table)

        story.append(Paragraph("Informações Técnicas", section_style))
        story.append(HRFlowable(width="100%", thickness=0.8, color=gray_line, spaceBefore=0, spaceAfter=8))
        technical_content = [[Paragraph(note, small_style)] for note in layout.technical_notes]
        technical_table = Table(
            [[Table([[""]], colWidths=[4]), Table(technical_content, colWidths=[doc.width - 28])]],
            colWidths=[8, doc.width - 8],
        )
        technical_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#CBD5E0")),
                    ("BACKGROUND", (1, 0), (1, 0), gray_soft),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        story.append(technical_table)

        doc.build(story)
        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'pdf')}\""
        return response

    story.append(Paragraph(DOCUMENT_TITLE, title_style))
    story.append(Paragraph("Cuidar Juntos", subtitle_style))
    summary_data = [
        [Paragraph(label, header_style), Paragraph(value, cell_style)]
        for label, value in meta.summary_rows()
    ]
    summary_table = Table(summary_data, colWidths=[doc.width * 0.32, doc.width * 0.68])
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), header_bg),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("GRID", (0, 0), (-1, -1), 0.5, grid_gray),
            ]
        )
    )
    story.append(summary_table)
    story.append(Spacer(1, 16))

    data: list[list[object]] = [[Paragraph(label, header_style) for _, label in columns]]
    data.extend(
        [
            [Paragraph(row.get(key, ""), cell_style) for key, _ in columns]
            for row in rows
        ]
    )

    col_units = _column_width_units(columns)
    unit_total = sum(col_units) or 1
    col_widths = [(doc.width * unit / unit_total) for unit in col_units]
    table = Table(data, repeatRows=1, colWidths=col_widths)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), header_bg),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("ALIGN", (0, 1), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 1), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("GRID", (0, 0), (-1, -1), 0.5, grid_gray),
            ]
        )
    )
    story.append(table)

    doc.build(story)
    response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'pdf')}\""
    return response


def _export_xlsx_inline(
    rows: list[dict[str, str]],
    meta: ExportMetadata,
    columns: Sequence[tuple[str, str]] = COLUMNS,
) -> HttpResponse:
    headers = [label for _, label in columns]
    column_widths = _column_width_units(columns)
    width = len(headers)

    all_rows: list[list[str]] = []
    all_rows.append(headers)
    all_rows.extend([[row.get(key, "") for key, _ in columns] for row in rows])

    last_col = _excel_column_letter(len(headers) - 1)
    last_row = len(all_rows)
    dimension = f"A1:{last_col}{last_row}"

    sheet_rows: list[str] = []
    for row_index, values in enumerate(all_rows, start=1):
        cells = []
        for col_index, value in enumerate(values):
            cell_ref = f"{_excel_column_letter(col_index)}{row_index}"
            safe_value = _xlsx_escape(value)
            cells.append(
                f'<c r="{cell_ref}" t="inlineStr"><is><t xml:space="preserve">{safe_value}</t></is></c>'
            )
        sheet_rows.append(f'<row r="{row_index}">{"".join(cells)}</row>')

    cols_xml = "".join(
        f'<col min="{idx}" max="{idx}" width="{column_widths[idx - 1]}" customWidth="1"/>'
        for idx in range(1, len(headers) + 1)
    )
    sheet_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n""" + (
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<dimension ref="{dimension}"/>'
        '<sheetViews><sheetView workbookViewId="0"/></sheetViews>'
        '<sheetFormatPr defaultRowHeight="15"/>'
        f"<cols>{cols_xml}</cols>"
        f"<sheetData>{''.join(sheet_rows)}</sheetData>"
        "</worksheet>"
    )

    created = timezone.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    with BytesIO() as buffer:
        with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(
                "[Content_Types].xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
                "<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
                "<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
                "<Override PartName=\"/xl/workbook.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml\"/>"
                "<Override PartName=\"/xl/worksheets/sheet1.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml\"/>"
                "<Override PartName=\"/docProps/core.xml\" ContentType=\"application/vnd.openxmlformats-package.core-properties+xml\"/>"
                "<Override PartName=\"/docProps/app.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.extended-properties+xml\"/>"
                "<Override PartName=\"/xl/styles.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml\"/>"
                "</Types>",
            )
            zf.writestr(
                "_rels/.rels",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
                "<Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" Target=\"xl/workbook.xml\"/>"
                "<Relationship Id=\"rId2\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties\" Target=\"docProps/app.xml\"/>"
                "<Relationship Id=\"rId3\" Type=\"http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties\" Target=\"docProps/core.xml\"/>"
                "</Relationships>",
            )
            zf.writestr(
                "docProps/app.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<Properties xmlns=\"http://schemas.openxmlformats.org/officeDocument/2006/extended-properties\" "
                "xmlns:vt=\"http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes\">"
                "<Application>Cuidar Juntos</Application>"
                "<TotalTime>0</TotalTime>"
                "</Properties>",
            )
            zf.writestr(
                "docProps/core.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<cp:coreProperties xmlns:cp=\"http://schemas.openxmlformats.org/package/2006/metadata/core-properties\" "
                "xmlns:dc=\"http://purl.org/dc/elements/1.1/\" "
                "xmlns:dcterms=\"http://purl.org/dc/terms/\" "
                "xmlns:dcmitype=\"http://purl.org/dc/dcmitype/\" xmlns:xsi=\"http://www.w3.org/2001/XMLSchema-instance\">"
                f"<dc:title>Registros exportados</dc:title><cp:category>{xml_escape(meta.period_label)}</cp:category>"
                f"<dcterms:created xsi:type=\"dcterms:W3CDTF\">{created}</dcterms:created>"
                "</cp:coreProperties>",
            )
            zf.writestr(
                "xl/workbook.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<workbook xmlns=\"http://schemas.openxmlformats.org/spreadsheetml/2006/main\" "
                "xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\">"
                "<bookViews><workbookView/></bookViews>"
                "<sheets><sheet name=\"Registros\" sheetId=\"1\" r:id=\"rId1\"/></sheets>"
                "</workbook>",
            )
            zf.writestr(
                "xl/_rels/workbook.xml.rels",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<Relationships xmlns=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\">"
                "<Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet\" Target=\"worksheets/sheet1.xml\"/>"
                "<Relationship Id=\"rId2\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles\" Target=\"styles.xml\"/>"
                "</Relationships>",
            )
            zf.writestr(
                "xl/styles.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<styleSheet xmlns=\"http://schemas.openxmlformats.org/spreadsheetml/2006/main\">"
                "<fonts count=\"1\"><font><name val=\"Calibri\"/><family val=\"2\"/><sz val=\"11\"/></font></fonts>"
                "<fills count=\"1\"><fill><patternFill patternType=\"none\"/></fill></fills>"
                "<borders count=\"1\"><border><left/><right/><top/><bottom/><diagonal/></border></borders>"
                "<cellStyleXfs count=\"1\"><xf numFmtId=\"0\" fontId=\"0\" fillId=\"0\" borderId=\"0\"/></cellStyleXfs>"
                "<cellXfs count=\"1\"><xf numFmtId=\"0\" fontId=\"0\" fillId=\"0\" borderId=\"0\" xfId=\"0\"/></cellXfs>"
                "<cellStyles count=\"1\"><cellStyle name=\"Normal\" xfId=\"0\" builtinId=\"0\"/></cellStyles>"
                "</styleSheet>",
            )
            zf.writestr("xl/worksheets/sheet1.xml", sheet_xml)

        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'xlsx')}\""
    return response


def _export_docx_inline(
    rows: list[dict[str, str]],
    meta: ExportMetadata,
    columns: Sequence[tuple[str, str]] = COLUMNS,
) -> HttpResponse:
    def paragraph(text: str) -> str:
        if not text:
            return "<w:p/>"
        escaped = xml_escape(text)
        escaped = escaped.replace("\n", "<w:br/>")
        return f"<w:p><w:r><w:t xml:space=\"preserve\">{escaped}</w:t></w:r></w:p>"

    table_rows = []
    headers = [label for _, label in columns]
    widths_twips = [px * 15 for px in _column_width_px(columns)]
    table_rows.append(_docx_table_row(headers, widths_twips, bold=True))
    for row in rows:
        values = [row.get(key, "") for key, _ in columns]
        table_rows.append(_docx_table_row(values, widths_twips))

    table_xml = (
        "<w:tbl>"
        "<w:tblPr>"
        "<w:tblStyle w:val=\"TableGrid\"/><w:tblW w:w=\"0\" w:type=\"auto\"/>"
        "<w:tblBorders>"
        "<w:top w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:left w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:bottom w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:right w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:insideH w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:insideV w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "</w:tblBorders>"
        "</w:tblPr>"
        "<w:tblGrid>"
        f"{''.join(f'<w:gridCol w:w=\"{width}\"/>' for width in widths_twips)}"
        "</w:tblGrid>"
        f"{''.join(table_rows)}"
        "</w:tbl>"
    )
    header_widths_twips = [2600, 6400]
    header_rows = [
        _docx_table_row([label, value], header_widths_twips, bold=[True, False])
        for label, value in meta.summary_rows()
    ]
    header_table_xml = (
        "<w:tbl>"
        "<w:tblPr>"
        "<w:tblStyle w:val=\"TableGrid\"/><w:tblW w:w=\"0\" w:type=\"auto\"/>"
        "<w:tblBorders>"
        "<w:top w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:left w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:bottom w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:right w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:insideH w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "<w:insideV w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
        "</w:tblBorders>"
        "</w:tblPr>"
        "<w:tblGrid>"
        f"{''.join(f'<w:gridCol w:w=\"{width}\"/>' for width in header_widths_twips)}"
        "</w:tblGrid>"
        f"{''.join(header_rows)}"
        "</w:tbl>"
    )

    document_xml = (
        """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
        "<w:document xmlns:wpc=\"http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas\" "
        "xmlns:mc=\"http://schemas.openxmlformats.org/markup-compatibility/2006\" "
        "xmlns:o=\"urn:schemas-microsoft-com:office:office\" "
        "xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\" "
        "xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\" "
        "xmlns:v=\"urn:schemas-microsoft-com:vml\" "
        "xmlns:wp14=\"http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing\" "
        "xmlns:wp=\"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing\" "
        "xmlns:w10=\"urn:schemas-microsoft-com:office:word\" "
        "xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" "
        "xmlns:w14=\"http://schemas.microsoft.com/office/word/2010/wordml\" "
        "xmlns:wpg=\"http://schemas.microsoft.com/office/word/2010/wordprocessingGroup\" "
        "xmlns:wpi=\"http://schemas.microsoft.com/office/word/2010/wordprocessingInk\" "
        "xmlns:wne=\"http://schemas.microsoft.com/office/word/2006/wordml\" "
        "xmlns:wps=\"http://schemas.microsoft.com/office/word/2010/wordprocessingShape\" mc:Ignorable=\"w14 wp14\">"
        "<w:body>"
        f"{paragraph('Cuidar Juntos')}"
        f"{paragraph(DOCUMENT_TITLE)}"
        f"{header_table_xml}"
        f"{paragraph('')}"
        f"{table_xml}"
        "<w:sectPr>"
        "<w:pgSz w:w=\"12240\" w:h=\"15840\"/><w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\"/>"
        "</w:sectPr>"
        "</w:body>"
        "</w:document>"
    )

    created = timezone.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    with BytesIO() as buffer:
        with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(
                "[Content_Types].xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
                "<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
                "<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
                "<Override PartName=\"/word/document.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>"
                "<Override PartName=\"/word/styles.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml\"/>"
                "<Override PartName=\"/docProps/core.xml\" ContentType=\"application/vnd.openxmlformats-package.core-properties+xml\"/>"
                "<Override PartName=\"/docProps/app.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.extended-properties+xml\"/>"
                "</Types>",
            )
            zf.writestr(
                "_rels/.rels",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
                "<Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" Target=\"word/document.xml\"/>"
                "<Relationship Id=\"rId2\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties\" Target=\"docProps/app.xml\"/>"
                "<Relationship Id=\"rId3\" Type=\"http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties\" Target=\"docProps/core.xml\"/>"
                "</Relationships>",
            )
            zf.writestr(
                "docProps/app.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<Properties xmlns=\"http://schemas.openxmlformats.org/officeDocument/2006/extended-properties\" "
                "xmlns:vt=\"http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes\">"
                "<Application>Cuidar Juntos</Application>"
                "<DocSecurity>0</DocSecurity>"
                "</Properties>",
            )
            zf.writestr(
                "docProps/core.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<cp:coreProperties xmlns:cp=\"http://schemas.openxmlformats.org/package/2006/metadata/core-properties\" "
                "xmlns:dc=\"http://purl.org/dc/elements/1.1/\" "
                "xmlns:dcterms=\"http://purl.org/dc/terms/\" "
                "xmlns:dcmitype=\"http://purl.org/dc/dcmitype/\" xmlns:xsi=\"http://www.w3.org/2001/XMLSchema-instance\">"
                "<dc:title>Registros exportados</dc:title>"
                f"<dc:subject>{xml_escape(meta.period_label)}</dc:subject>"
                f"<dcterms:created xsi:type=\"dcterms:W3CDTF\">{created}</dcterms:created>"
                "</cp:coreProperties>",
            )
            zf.writestr("word/document.xml", document_xml)
            zf.writestr(
                "word/styles.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>"""
                "<w:styles xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
                "<w:style w:type=\"paragraph\" w:default=\"1\" w:styleId=\"Normal\">"
                "<w:name w:val=\"Normal\"/><w:qFormat/>"
                "<w:rPr><w:lang w:val=\"pt-BR\"/></w:rPr>"
                "</w:style>"
                "<w:style w:type=\"table\" w:styleId=\"TableGrid\">"
                "<w:name w:val=\"Table Grid\"/>"
                "<w:tblPr><w:tblBorders>"
                "<w:top w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                "<w:left w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                "<w:bottom w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                "<w:right w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                "<w:insideH w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                "<w:insideV w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"auto\"/>"
                "</w:tblBorders></w:tblPr>"
                "</w:style>"
                "</w:styles>",
            )

        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'docx')}\""
    return response


def _export_pdf_inline(
    rows: list[dict[str, str]],
    meta: ExportMetadata,
    columns: Sequence[tuple[str, str]] = COLUMNS,
) -> HttpResponse:
    lines: list[str] = [DOCUMENT_TITLE, "Cuidar Juntos", ""]
    lines.extend(f"{label}: {value}" for label, value in meta.summary_rows())
    lines.append("")
    header = " | ".join(label for _, label in columns)
    lines.append(header)
    lines.append("-" * len(header))

    for row in rows:
        text = " | ".join((row.get(key, "") or "-") for key, _ in columns)
        wrapped = textwrap.wrap(text, width=100) or [""]
        lines.extend(wrapped)
        lines.append("")

    if not rows:
        lines.append("Nenhum registro encontrado para os filtros selecionados.")

    pdf_bytes = _build_simple_pdf(lines)
    response = HttpResponse(pdf_bytes, content_type="application/pdf")
    response["Content-Disposition"] = f"attachment; filename=\"{_default_filename(meta, 'pdf')}\""
    return response


def _build_simple_pdf(lines: list[str]) -> bytes:
    pages: list[list[str]] = []
    page: list[str] = []
    max_lines = 45
    for line in lines:
        if len(page) >= max_lines:
            pages.append(page)
            page = []
        page.append(line)
    if page:
        pages.append(page)
    if not pages:
        pages = [[""]]

    num_pages = len(pages)
    total_objects = 2 + num_pages * 2 + 1  # catalog + pages + (page+content)*n + font
    objects: list[bytes | None] = [None] * (total_objects + 1)

    catalog_obj = 1
    pages_obj = 2
    page_objs = list(range(3, 3 + num_pages))
    content_objs = list(range(3 + num_pages, 3 + (num_pages * 2)))
    font_obj = total_objects

    objects[font_obj] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"

    for idx, page_lines in enumerate(pages):
        content_stream = _pdf_page_stream(page_lines)
        content_obj = content_objs[idx]
        objects[content_obj] = content_stream
        page_obj = page_objs[idx]
        objects[page_obj] = (
            f"<< /Type /Page /Parent {pages_obj} 0 R /MediaBox [0 0 595 842] "
            f"/Contents {content_obj} 0 R /Resources << /Font << /F1 {font_obj} 0 R >> >> >>"
        ).encode("utf-8")

    kids = " ".join(f"{num} 0 R" for num in page_objs)
    objects[pages_obj] = f"<< /Type /Pages /Count {num_pages} /Kids [{kids}] >>".encode("utf-8")
    objects[catalog_obj] = f"<< /Type /Catalog /Pages {pages_obj} 0 R >>".encode("utf-8")

    buffer = BytesIO()
    buffer.write(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")

    offsets = [0] * (total_objects + 1)
    for obj_num in range(1, total_objects + 1):
        offsets[obj_num] = buffer.tell()
        buffer.write(f"{obj_num} 0 obj\n".encode("utf-8"))
        data = objects[obj_num]
        if isinstance(data, str):
            data = data.encode("utf-8")
        buffer.write(data or b"")
        buffer.write(b"\nendobj\n")

    xref_pos = buffer.tell()
    buffer.write(f"xref\n0 {total_objects + 1}\n".encode("utf-8"))
    buffer.write(b"0000000000 65535 f \n")
    for obj_num in range(1, total_objects + 1):
        buffer.write(f"{offsets[obj_num]:010d} 00000 n \n".encode("utf-8"))
    buffer.write(
        f"trailer\n<< /Size {total_objects + 1} /Root {catalog_obj} 0 R >>\nstartxref\n{xref_pos}\n%%EOF".encode(
            "utf-8"
        )
    )
    return buffer.getvalue()


def _pdf_page_stream(lines: list[str]) -> bytes:
    ops = ["BT", "/F1 11 Tf", "14 TL", "50 800 Td"]
    for line in lines:
        text = _pdf_escape(line)
        ops.append(f"({text}) Tj")
        ops.append("T*")
    ops.append("ET")
    content = "\n".join(ops).encode("utf-8")
    return b"<< /Length " + str(len(content)).encode("utf-8") + b" >>\nstream\n" + content + b"\nendstream"


def _pdf_escape(value: str) -> str:
    return (
        (value or "")
        .replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
        .replace("\r", " ")
        .replace("\n", " ")
    )


def _docx_table_row(
    values: list[str],
    widths_twips: Sequence[int],
    bold: bool | Sequence[bool] = False,
) -> str:
    if isinstance(bold, bool):
        bold_flags = [bold] * len(values)
    else:
        bold_flags = list(bold) + [False] * max(0, len(values) - len(bold))  # type: ignore[arg-type]
        bold_flags = bold_flags[: len(values)]

    cells = []
    for idx, value in enumerate(values):
        escaped = xml_escape(value or "")
        run_props = "<w:rPr><w:b/></w:rPr>" if bold_flags[idx] else ""
        width = widths_twips[idx] if idx < len(widths_twips) else widths_twips[-1]
        cells.append(
            "<w:tc><w:tcPr>"
            f"<w:tcW w:w=\"{width}\" w:type=\"dxa\"/>"
            "</w:tcPr>"
            "<w:p><w:pPr><w:jc w:val=\"center\"/></w:pPr>"
            f"<w:r>{run_props}<w:t xml:space=\"preserve\">{escaped}</w:t></w:r></w:p></w:tc>"
        )
    return f"<w:tr>{''.join(cells)}</w:tr>"


def _xlsx_escape(value: str) -> str:
    return (xml_escape(_xlsx_sanitize(value))).replace("\n", "&#10;")


_XLSX_ILLEGAL_CHARS_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def _xlsx_sanitize(value: str | None) -> str:
    if value is None:
        return ""
    if not isinstance(value, str):
        value = str(value)
    return _XLSX_ILLEGAL_CHARS_RE.sub("", value)


def _excel_column_letter(index: int) -> str:
    result = ""
    while index >= 0:
        index, remainder = divmod(index, 26)
        result = chr(65 + remainder) + result
        index -= 1
    return result


def _render_sleep_chart_svg(sessions: Sequence[dict[str, object]], meta: ExportMetadata) -> str:
    width = 960
    height = 620
    margin_left = 70
    margin_right = 24
    margin_top = 70
    margin_bottom = 150

    chart_width = width - margin_left - margin_right
    chart_height = height - margin_top - margin_bottom

    if not sessions:
        return (
            f"<svg xmlns=\"http://www.w3.org/2000/svg\" width=\"{width}\" height=\"{height}\">"
            "<rect width=\"100%\" height=\"100%\" fill=\"#FFFFFF\"/>"
            f"<text x=\"{width/2}\" y=\"{height/2}\" text-anchor=\"middle\" "
            "font-family=\"Arial\" font-size=\"16\" fill=\"#334155\">"
            "Nenhum dado de sono para o período selecionado."
            "</text></svg>"
        )

    date_map = {session["date"]: session["date_label"] for session in sessions}
    dates = sorted(date_map.keys())
    max_hours = max(float(session["hours"]) for session in sessions)
    max_y = max(1, math.ceil(max_hours))

    if max_y <= 6:
        y_step = 1
    elif max_y <= 12:
        y_step = 2
    else:
        y_step = 3

    by_patient: dict[str, list[dict[str, object]]] = {}
    for session in sessions:
        patient = str(session["patient"])
        by_patient.setdefault(patient, []).append(session)

    colors = ["#2563EB", "#10B981", "#F59E0B", "#EF4444", "#8B5CF6", "#14B8A6"]

    def y_for(hours: float) -> float:
        return margin_top + chart_height * (1 - (hours / max_y))

    svg_parts = [
        f"<svg xmlns=\"http://www.w3.org/2000/svg\" width=\"{width}\" height=\"{height}\">",
        "<rect width=\"100%\" height=\"100%\" fill=\"#FFFFFF\"/>",
        f"<text x=\"{margin_left}\" y=\"32\" font-family=\"Arial\" font-size=\"18\" "
        "fill=\"#0F172A\" font-weight=\"bold\">Gráfico de Sono</text>",
        f"<text x=\"{margin_left}\" y=\"52\" font-family=\"Arial\" font-size=\"12\" "
        "fill=\"#475569\">"
        f"{xml_escape(meta.describe())}"
        "</text>",
    ]

    for tick in range(0, max_y + 1, y_step):
        y = y_for(float(tick))
        svg_parts.append(
            f"<line x1=\"{margin_left}\" y1=\"{y}\" x2=\"{width - margin_right}\" "
            f"y2=\"{y}\" stroke=\"#E2E8F0\" stroke-width=\"1\"/>"
        )
        svg_parts.append(
            f"<text x=\"{margin_left - 8}\" y=\"{y + 4}\" text-anchor=\"end\" "
            "font-family=\"Arial\" font-size=\"11\" fill=\"#64748B\">"
            f"{tick}</text>"
        )

    svg_parts.append(
        f"<line x1=\"{margin_left}\" y1=\"{margin_top}\" x2=\"{margin_left}\" "
        f"y2=\"{height - margin_bottom}\" stroke=\"#0F172A\" stroke-width=\"1\"/>"
    )
    svg_parts.append(
        f"<line x1=\"{margin_left}\" y1=\"{height - margin_bottom}\" "
        f"x2=\"{width - margin_right}\" y2=\"{height - margin_bottom}\" "
        "stroke=\"#0F172A\" stroke-width=\"1\"/>"
    )

    date_positions: dict[object, float] = {}
    group_width = chart_width / max(1, len(dates))
    for idx, label in enumerate(dates):
        group_start = margin_left + idx * group_width
        date_positions[label] = group_start + group_width / 2

    max_labels = 10
    x_label_step = max(1, math.ceil(len(dates) / max_labels))
    for idx, label in enumerate(dates):
        if idx % x_label_step != 0:
            continue
        x = date_positions[label]
        svg_parts.append(
            f"<text x=\"{x}\" y=\"{height - margin_bottom + 18}\" text-anchor=\"middle\" "
            "font-family=\"Arial\" font-size=\"10\" fill=\"#64748B\">"
            f"{xml_escape(str(date_map[label]))}</text>"
        )

    svg_parts.append(
        f"<text x=\"{margin_left}\" y=\"{height - margin_bottom + 30}\" font-family=\"Arial\" "
        "font-size=\"12\" fill=\"#475569\">Data do registro (Dormiu)</text>"
    )
    svg_parts.append(
        f"<text x=\"20\" y=\"{margin_top - 12}\" font-family=\"Arial\" "
        "font-size=\"12\" fill=\"#475569\">Horas de sono</text>"
    )

    legend_y = margin_top - 8
    legend_x = width - margin_right - 150
    patient_names = [name for name, _ in sorted(by_patient.items())]
    patient_count = max(1, len(patient_names))
    bar_gap = 6
    inner_padding = 0.12
    bar_slot_width = group_width * (1 - inner_padding * 2)
    bar_width = max(6, (bar_slot_width - (patient_count - 1) * bar_gap) / patient_count)

    hours_by_date_patient: dict[object, dict[str, float]] = {label: {} for label in dates}
    for patient, items in by_patient.items():
        for item in items:
            hours_by_date_patient[item["date"]][patient] = float(item["hours"])

    for patient_index, patient in enumerate(patient_names):
        color = colors[patient_index % len(colors)]
        y = legend_y + patient_index * 16
        svg_parts.append(
            f"<rect x=\"{legend_x}\" y=\"{y - 10}\" width=\"10\" height=\"10\" "
            f"fill=\"{color}\"/>"
        )
        svg_parts.append(
            f"<text x=\"{legend_x + 14}\" y=\"{y - 1}\" font-family=\"Arial\" "
            "font-size=\"10\" fill=\"#475569\">"
            f"{xml_escape(patient)}</text>"
        )

        for date_index, label in enumerate(dates):
            hours = hours_by_date_patient[label].get(patient)
            if hours is None:
                continue
            bar_height = chart_height * (hours / max_y)
            x = margin_left + date_index * group_width + (group_width - bar_slot_width) / 2
            x += patient_index * (bar_width + bar_gap)
            y_top = margin_top + chart_height - bar_height
            svg_parts.append(
                f"<rect x=\"{x:.2f}\" y=\"{y_top:.2f}\" width=\"{bar_width:.2f}\" "
                f"height=\"{bar_height:.2f}\" rx=\"3\" ry=\"3\" fill=\"{color}\"/>"
            )

    svg_parts.append("</svg>")
    svg = "".join(svg_parts)

    total_hours = sum(float(session["hours"]) for session in sessions)
    avg_hours = total_hours / len(sessions)
    max_session = max(sessions, key=lambda row: float(row["hours"]))
    min_session = min(sessions, key=lambda row: float(row["hours"]))
    summary_lines = [
        f"Media de horas dormidas: {avg_hours:.2f} h",
        f"Maior sono: {max_session['date_label']} - {float(max_session['hours']):.2f} h ({max_session['patient']})",
        f"Menor sono: {min_session['date_label']} - {float(min_session['hours']):.2f} h ({min_session['patient']})",
    ]
    summary_y_start = height - margin_bottom + 60
    summary_svg = []
    for idx, line in enumerate(summary_lines):
        y = summary_y_start + idx * 18
        summary_svg.append(
            f"<text x=\"{margin_left}\" y=\"{y}\" font-family=\"Arial\" "
            "font-size=\"12\" fill=\"#0F172A\">"
            f"{xml_escape(line)}</text>"
        )
    svg = svg.replace("</svg>", "".join(summary_svg) + "</svg>")

def export_sleep_chart(sessions: Sequence[dict[str, object]], meta: ExportMetadata) -> HttpResponse:
    svg = _render_sleep_chart_svg(sessions, meta)
    response = HttpResponse(svg, content_type="image/svg+xml")
    response["Content-Disposition"] = (
        f"attachment; filename=sono_grafico_{meta.range_slug}_{meta.patient_slug}.svg"
    )
    return response




EXPORTERS = {
    "csv": export_as_csv,
    "xlsx": export_as_xlsx,
    "docx": export_as_docx,
    "pdf": export_as_pdf,
    "sleep_chart": export_sleep_chart,
}
